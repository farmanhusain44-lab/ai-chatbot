from flask import Flask, request, jsonify, send_from_directory, Response, redirect
from flask_cors import CORS
import anthropic
import json
import os
import re
import shutil
import tempfile
import time
import uuid
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from datetime import datetime, timedelta
from database import init_db, create_client, get_client, get_client_by_access_code, get_all_clients, add_document, get_documents, get_client_context, increment_message_count
from media_editor import edit_media, kind_from_path, MediaEditorError, SUPPORTED_OPS_HINT
from instruction_parser import parse_instruction as parse_edit_instruction
from replicate_ops import ai_transcribe_to_srt, ai_transcribe_words, ReplicateUnavailable
import subprocess
from twilio.rest import Client as TwilioClient
from twilio.twiml.messaging_response import MessagingResponse
from langdetect import detect
from langdetect.lang_detect_exception import LangDetectException
import requests
import logging
from io import BytesIO

try:
    import openai
except ImportError:
    openai = None

# Document extraction libraries
try:
    import PyPDF2
except ImportError:
    PyPDF2 = None
try:
    from docx import Document
except ImportError:
    Document = None

# Configure logging so errors appear in Gunicorn/Railway logs
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s"
)
logger = logging.getLogger(__name__)

app = Flask(__name__, static_folder='static')
CORS(app)
port = int(os.environ.get("PORT", 8080))

# Initialize database
init_db()

# In-memory document knowledge base
DOCUMENT_CHUNKS = []
MAX_DOCUMENT_CHUNKS = 2000
CHUNK_SIZE = 800
CHUNK_OVERLAP = 100

# In-memory lead storage (in production, use database)
LEADS = []

# In-memory clients DB (payment records)
clients_db = []


def normalize_arabic(text):
    """Normalize Arabic text for better matching across dialects and diacritics."""
    if not text:
        return text
    # Remove tashkeel (diacritics)
    text = re.sub(r'[\u064B-\u065F\u0670\u0640]', '', text)
    # Normalize alef variants
    text = re.sub(r'[\u0622\u0623\u0625]', '\u0627', text)
    # Normalize yaa and alif maqsura
    text = text.replace('\u0649', '\u064A')
    return text


def extract_text_from_file(file_input, filename):
    """Extract text from PDF, DOCX, or TXT files. file_input can be a file object or a path string."""
    ext = os.path.splitext(filename.lower())[1]
    try:
        if ext == '.pdf' and PyPDF2:
            if isinstance(file_input, str):
                reader = PyPDF2.PdfReader(file_input)
            else:
                reader = PyPDF2.PdfReader(BytesIO(file_input.read()))
            parts = []
            for page in reader.pages:
                try:
                    parts.append(page.extract_text() or '')
                except Exception:
                    pass
            return '\n'.join(parts)
        elif ext in ('.docx', '.doc') and Document:
            if isinstance(file_input, str):
                with open(file_input, 'rb') as f:
                    doc = Document(BytesIO(f.read()))
            else:
                doc = Document(BytesIO(file_input.read()))
            return '\n'.join(p.text for p in doc.paragraphs if p.text)
        elif ext in ('.txt', '.md', '.csv', '.json'):
            if isinstance(file_input, str):
                with open(file_input, 'r', encoding='utf-8', errors='ignore') as f:
                    return f.read()
            return file_input.read().decode('utf-8', errors='ignore')
        else:
            if isinstance(file_input, str):
                with open(file_input, 'r', encoding='utf-8', errors='ignore') as f:
                    return f.read()
            return file_input.read().decode('utf-8', errors='ignore')
    except Exception as e:
        logger.error("Document extraction failed for %s: %s", filename, e)
        return ""


def chunk_text(text, chunk_size=CHUNK_SIZE, overlap=CHUNK_OVERLAP):
    """Split text into overlapping chunks."""
    text = re.sub(r'\s+', ' ', text).strip()
    if not text:
        return []
    chunks = []
    start = 0
    while start < len(text):
        end = min(start + chunk_size, len(text))
        # Try to break at a sentence or space
        if end < len(text):
            for sep in ['. ', '? ', '! ', '\n', ' ']:
                pos = text.rfind(sep, start, end)
                if pos > start:
                    end = pos + len(sep)
                    break
        chunks.append(text[start:end].strip())
        start = end - overlap if end < len(text) else end
    return chunks


def embed_texts(texts):
    """Get OpenAI embeddings for a list of texts. Falls back to None if unavailable."""
    oc = get_openai_client()
    if not oc or not texts:
        return None
    try:
        response = oc.embeddings.create(
            model="text-embedding-3-small",
            input=texts
        )
        return [item.embedding for item in response.data]
    except Exception as e:
        logger.error("OpenAI embedding failed: %s", e)
        return None


def cosine_similarity(a, b):
    import numpy as np
    return float(np.dot(a, b) / (np.linalg.norm(a) * np.linalg.norm(b)))


def add_document_to_knowledge_base(filename, text):
    """Chunk, embed, and store document text."""
    global DOCUMENT_CHUNKS
    chunks = chunk_text(text)
    valid_chunks = [chunk for chunk in chunks if chunk]
    if not valid_chunks:
        return 0

    embeddings = embed_texts(valid_chunks)
    for i, chunk in enumerate(valid_chunks):
        item = {"source": filename, "text": chunk, "norm": normalize_arabic(chunk.lower())}
        if embeddings:
            item["embedding"] = embeddings[i]
        DOCUMENT_CHUNKS.append(item)

    # Keep within memory limit
    if len(DOCUMENT_CHUNKS) > MAX_DOCUMENT_CHUNKS:
        DOCUMENT_CHUNKS = DOCUMENT_CHUNKS[-MAX_DOCUMENT_CHUNKS:]
    logger.info("Added document %s: %d chunks, total %d chunks, embeddings=%s", filename, len(valid_chunks), len(DOCUMENT_CHUNKS), bool(embeddings))
    return len(valid_chunks)


def keyword_search_context(query, top_k=3):
    """Keyword search as fallback."""
    if not DOCUMENT_CHUNKS:
        return ""
    query_norm = normalize_arabic(query.lower())
    query_words = set(re.findall(r"[\w'\u0600-\u06FF]+", query_norm))
    # Stop words for English and Arabic
    stop_words = {
        'the', 'is', 'a', 'an', 'and', 'or', 'to', 'of', 'in', 'for', 'on', 'with', 'what', 'how', 'who', 'where', 'when', 'why', 'me', 'my', 'your', 'this', 'that', 'are', 'do', 'does', 'can', 'you', 'was', 'were', 'did', 'will',
        'هذا', 'هذه', 'التي', 'الذي', 'من', 'في', 'على', 'إلى', 'عن', 'مع', 'كان', 'أن', 'أو', 'لم', 'قد', 'ما', 'كل', 'بعد', 'قبل', 'أي', 'هل', 'كيف', 'أين', 'متى', 'لماذا', 'لمن', 'هو', 'هي', 'هم'
    }
    query_words = query_words - stop_words
    if not query_words:
        return ""
    scored = []
    for chunk in DOCUMENT_CHUNKS:
        chunk_words = set(re.findall(r"[\w'\u0600-\u06FF]+", chunk.get("norm", chunk["text"].lower())))
        score = len(query_words & chunk_words)
        if score > 0:
            scored.append((score, chunk["text"]))
    scored.sort(reverse=True, key=lambda x: x[0])
    selected = []
    total_len = 0
    for score, text in scored[:top_k]:
        if total_len + len(text) > 2500:
            break
        selected.append(text)
        total_len += len(text)
    return "\n\n".join(selected)


def get_relevant_context(query, top_k=3):
    """Find relevant document chunks using embeddings if available, else keyword search."""
    if not DOCUMENT_CHUNKS:
        return ""

    # If we have embeddings, do vector search
    if get_openai_client() and any("embedding" in c for c in DOCUMENT_CHUNKS):
        query_embedding = embed_texts([query])
        if query_embedding and query_embedding[0]:
            import numpy as np
            q_vec = np.array(query_embedding[0])
            scored = []
            for chunk in DOCUMENT_CHUNKS:
                if "embedding" not in chunk:
                    continue
                c_vec = np.array(chunk["embedding"])
                score = cosine_similarity(q_vec, c_vec)
                scored.append((score, chunk["text"]))
            scored.sort(reverse=True, key=lambda x: x[0])
            selected = []
            total_len = 0
            for score, text in scored[:top_k]:
                if total_len + len(text) > 2500:
                    break
                selected.append(text)
                total_len += len(text)
            if selected:
                return "\n\n".join(selected)

    # Fallback to keyword search
    return keyword_search_context(query, top_k)


@app.route('/public/<path:filename>')
def serve_public(filename):
    return send_from_directory('public', filename)

@app.route('/sitemap.xml')
def sitemap():
    return send_from_directory('static', 'sitemap.xml', mimetype='application/xml')

@app.route('/robots.txt')
def robots():
    return send_from_directory('static', 'robots.txt', mimetype='text/plain')

# Warn at startup if the API key is missing so it shows up in deploy logs
api_key = os.environ.get("ANTHROPIC_API_KEY")
if not api_key:
    logger.warning("ANTHROPIC_API_KEY environment variable is not set — API calls will fail")

_anthropic_client = None
def get_anthropic_client():
    global _anthropic_client
    if _anthropic_client is None and api_key:
        import anthropic
        _anthropic_client = anthropic.Anthropic(api_key=api_key)
    return _anthropic_client

# Optional OpenAI client for document embeddings (vector search)
openai_api_key = os.environ.get("OPENAI_API_KEY")
_openai_client = None
def get_openai_client():
    global _openai_client
    if _openai_client is None and openai and openai_api_key:
        try:
            _openai_client = openai.OpenAI(api_key=openai_api_key)
        except Exception as e:
            logger.warning("OpenAI client failed to initialize: %s", e)
    return _openai_client

twilio_sid = os.environ.get("TWILIO_ACCOUNT_SID")
twilio_token = os.environ.get("TWILIO_AUTH_TOKEN")
twilio_from = os.environ.get("TWILIO_WHATSAPP_NUMBER")
twilio = None
if twilio_sid and twilio_token:
    twilio = TwilioClient(twilio_sid, twilio_token)

elevenlabs_key = os.environ.get("ELEVENLABS_API_KEY")

LANGUAGE_NAMES = {
    "en": "English",
    "hi": "Hindi",
    "ur": "Urdu",
    "ar": "Arabic",
    "bn": "Bangla",
    "pa": "Punjabi",
    "ta": "Tamil",
    "te": "Telugu",
    "ml": "Malayalam",
    "kn": "Kannada",
    "mr": "Marathi",
    "gu": "Gujarati",
    "or": "Odia",
    "as": "Assamese",
    "es": "Spanish",
    "fr": "French",
    "de": "German",
    "zh": "Chinese",
    "ja": "Japanese",
    "ko": "Korean",
    "ru": "Russian",
    "pt": "Portuguese",
    "it": "Italian",
    "tr": "Turkish"
}

# Male voice IDs for ElevenLabs (Adam is a clear male voice that works across languages)
# All languages use the same multilingual voice; ElevenLabs model detects language automatically.
ELEVENLABS_VOICES = {k: "pNInz6obpgDQGcFmaJgB" for k in LANGUAGE_NAMES}

def detect_language(text):
    # Script-based detection is more reliable for non-Latin scripts
    def has_chars(start, end):
        return any(start <= ord(c) <= end for c in text)

    # Urdu-specific characters (Arabic script with Persian/Urdu letters)
    if has_chars(0x067E, 0x067E) or has_chars(0x0686, 0x0686) or has_chars(0x0688, 0x0688) \
            or has_chars(0x0691, 0x0691) or has_chars(0x06A9, 0x06A9) or has_chars(0x06AF, 0x06AF) \
            or has_chars(0x06BA, 0x06BA) or has_chars(0x06CC, 0x06CC) or has_chars(0x06D2, 0x06D2) \
            or has_chars(0x06C1, 0x06C1) or has_chars(0x06BE, 0x06BE):
        return "ur"
    # Arabic/Persian script (without Urdu-specific characters)
    if has_chars(0x0600, 0x06FF) or has_chars(0x0750, 0x077F) or has_chars(0x08A0, 0x08FF):
        return "ar"
    if has_chars(0x0900, 0x097F):
        return "hi"  # Devanagari (Hindi, Marathi, Nepali)
    if has_chars(0x0980, 0x09FF):
        return "bn"  # Bengali
    if has_chars(0x0A00, 0x0A7F):
        return "pa"  # Gurmukhi (Punjabi)
    if has_chars(0x0A80, 0x0AFF):
        return "gu"  # Gujarati
    if has_chars(0x0B00, 0x0B7F):
        return "or"  # Oriya/Odia
    if has_chars(0x0B80, 0x0BFF):
        return "ta"  # Tamil
    if has_chars(0x0C00, 0x0C7F):
        return "te"  # Telugu
    if has_chars(0x0C80, 0x0CFF):
        return "kn"  # Kannada
    if has_chars(0x0D00, 0x0D7F):
        return "ml"  # Malayalam
    if has_chars(0x4E00, 0x9FFF):
        return "zh"  # Chinese
    if has_chars(0x3040, 0x309F) or has_chars(0x30A0, 0x30FF):
        return "ja"  # Japanese
    if has_chars(0xAC00, 0xD7AF) or has_chars(0x1100, 0x11FF):
        return "ko"  # Korean
    if has_chars(0x0400, 0x04FF):
        return "ru"  # Cyrillic

    # Roman Hindi detection (speech transcribed in English letters)
    roman_hindi_words = {
        "kya", "kaise", "kaun", "kahan", "kab", "kyun", "kitna", "kaunsa",
        "main", "tum", "aap", "woh", "yeh", "hum", "sab", "log",
        "hoon", "ho", "hai", "hain", "tha", "thi", "the", "raha", "rahi", "kar", "kiya", "gaya", "diya",
        "accha", "theek", "nahi", "bilkul", "bahut", "thoda", "zyada", "kam", "achha",
        "bhi", "lekin", "kyunki", "agar", "toh", "ya", "aur", "par", "se", "ko", "mein", "pe", "tak",
        "shukriya", "dhanyawad", "namaste", "bhai", "yaar", "chal", "karo", "dekho"
    }
    words = set(re.findall(r"\b[a-z]+\b", text.lower()))
    if len(words.intersection(roman_hindi_words)) >= 2:
        return "hi"

    try:
        lang = detect(text)
        if lang in LANGUAGE_NAMES:
            return lang
        # Fallback mappings for close languages
        if lang == "fa":
            return "ur"
        return "en"
    except LangDetectException:
        return "en"

def is_group_message(sender):
    # Twilio individual: whatsapp:+1234567890
    # Twilio group: whatsapp:1203630... (no +)
    if not sender:
        return True
    number = sender.replace("whatsapp:", "")
    return not number.startswith("+")

def send_whatsapp_reply(to, reply):
    if not twilio or not twilio_from:
        return False
    try:
        twilio.messages.create(
            from_=f"whatsapp:{twilio_from}",
            body=reply,
            to=to
        )
        return True
    except Exception as e:
        print(f"WhatsApp send error: {e}")
        return False

def get_ai_reply(message, language, timezone=None, history=None, context=None, region=None):
    # If history is provided, it already contains the current user message as the last item.
    messages = history if history else [{"role": "user", "content": message}]
    if context and messages and messages[-1]["role"] == "user":
        lang_name = LANGUAGE_NAMES.get(language, language)
        context_prompt = (
            f"You must reply in {lang_name}. "
            "Use the following document context to answer the question. "
            "If the answer is not in the context, say you don't know. "
            "Do not make up information.\n\n"
            f"Context:\n{context}\n\n"
            "Question:\n"
        )
        messages[-1]["content"] = context_prompt + messages[-1]["content"]
    anthropic_client = get_anthropic_client()
    if not anthropic_client:
        return "I'm sorry, the AI service is not configured right now."
    response = anthropic_client.messages.create(
        model="claude-haiku-4-5-20251001",
        max_tokens=600,
        system=get_system_prompt(language, timezone, region),
        messages=messages
    )
    if not response.content:
        logger.warning("Anthropic response had no content blocks (stop_reason=%s)", response.stop_reason)
        return "I'm sorry, I couldn't generate a reply for that. Could you rephrase your message?"
    return response.content[0].text

REGION_PRICING = {
    "india":       "Starter ₹100/month (500 chats), Professional ₹150/month (3,000 chats + WhatsApp), Enterprise ₹200/month (unlimited).",
    "pakistan":      "Starter Rs 300/month, Professional Rs 500/month, Enterprise Rs 700/month.",
    "bangladesh":    "Starter ৳150/month, Professional ৳220/month, Enterprise ৳300/month.",
    "srilanka":      "Starter රු 350/month, Professional රු 500/month, Enterprise රු 700/month.",
    "nepal":         "Starter रु 160/month, Professional रु 240/month, Enterprise रु 320/month.",
    "bhutan":        "Starter Nu. 100/month, Professional Nu. 150/month, Enterprise Nu. 200/month.",
    "uae":           "Starter AED 5/month, Professional AED 7/month, Enterprise AED 10/month.",
    "qatar":         "Starter QAR 5/month, Professional QAR 7/month, Enterprise QAR 10/month.",
    "kuwait":        "Starter KWD 0.5/month, Professional KWD 0.7/month, Enterprise KWD 1/month.",
    "bahrain":       "Starter BHD 0.5/month, Professional BHD 0.7/month, Enterprise BHD 1/month.",
    "oman":          "Starter OMR 0.5/month, Professional OMR 0.7/month, Enterprise OMR 1/month.",
    "jordan":        "Starter JOD 1/month, Professional JOD 1.5/month, Enterprise JOD 2/month.",
    "iraq":          "Starter IQD 1500/month, Professional IQD 2200/month, Enterprise IQD 3000/month.",
    "lebanon":       "Starter $1/month, Professional $2/month, Enterprise $3/month.",
    "egypt":         "Starter EGP 60/month, Professional EGP 90/month, Enterprise EGP 120/month.",
    "china":         "Starter ¥10/month, Professional ¥15/month, Enterprise ¥20/month.",
    "myanmar":       "Starter K 2,500/month, Professional K 4,000/month, Enterprise K 5,500/month.",
    "indonesia":     "Starter Rp 20,000/month, Professional Rp 30,000/month, Enterprise Rp 40,000/month.",
    "philippines":   "Starter ₱70/month, Professional ₱100/month, Enterprise ₱130/month.",
    "vietnam":       "Starter ₫30,000/month, Professional ₫45,000/month, Enterprise ₫60,000/month.",
    "malaysia":      "Starter RM 5/month, Professional RM 8/month, Enterprise RM 10/month.",
    "thailand":      "Starter ฿40/month, Professional ฿60/month, Enterprise ฿80/month.",
    "nigeria":       "Starter ₦2,000/month, Professional ₦3,000/month, Enterprise ₦4,000/month.",
    "kenya":         "Starter KSh 150/month, Professional KSh 230/month, Enterprise KSh 300/month.",
    "southafrica":   "Starter R 25/month, Professional R 35/month, Enterprise R 50/month.",
    "ghana":         "Starter GH₵ 20/month, Professional GH₵ 30/month, Enterprise GH₵ 40/month.",
    "uganda":        "Starter USh 5,000/month, Professional USh 7,000/month, Enterprise USh 10,000/month.",
    "tanzania":      "Starter TSh 3,000/month, Professional TSh 5,000/month, Enterprise TSh 7,000/month.",
    "rwanda":        "Starter RF 1,500/month, Professional RF 2,500/month, Enterprise RF 3,500/month.",
    "mexico":        "Starter MXN $25/month, Professional MXN $40/month, Enterprise MXN $50/month.",
    "brazil":        "Starter R$ 6/month, Professional R$ 10/month, Enterprise R$ 12/month.",
    "colombia":      "Starter COP $5,000/month, Professional COP $7,000/month, Enterprise COP $10,000/month.",
    "argentina":     "Starter ARS $1,200/month, Professional ARS $2,000/month, Enterprise ARS $2,500/month.",
    "peru":          "Starter S/ 5/month, Professional S/ 8/month, Enterprise S/ 10/month.",
    "chile":         "Starter CLP $1,200/month, Professional CLP $2,000/month, Enterprise CLP $2,500/month.",
}

def get_system_prompt(language="en", timezone=None, region=None):
    try:
        from zoneinfo import ZoneInfo
        tz = ZoneInfo(timezone) if timezone else None
    except Exception:
        tz = None
    now = datetime.now(tz) if tz else datetime.now()
    now_str = now.strftime("%A, %B %d, %Y at %I:%M %p %Z")
    region_key = (region or "").strip().lower()
    region_pricing = REGION_PRICING.get(region_key)
    region_line = ""
    if region_pricing:
        region_line = (
            f"The visitor is on the BotifyAI {region_key.title()} page — only mention pricing for {region_key.title()}: "
            f"{region_pricing} Do NOT mix in pricing from other regions unless the visitor explicitly asks about another country. "
        )
    return (
        "You are BotifyAI Assistant — a smart, warm, professional, and highly capable multilingual AI assistant. "
        + region_line +
        "ANSWER LENGTH — this is a strict rule the user has asked for: reply in 1-2 short sentences. Never more than 2 sentences unless the user EXPLICITLY says 'explain in detail' or 'tell me more'. No paragraphs, no bullet lists, no headings, no extra facts the user didn't ask about. Long replies annoy customers — think of it like a WhatsApp reply, not an essay. "
        "You can help with anything the user asks, not just BotifyAI topics. Answer the specific thing asked — nothing extra. If they ask 'what's 2+2', say '4.' — do not explain arithmetic. If they ask 'how many languages in the world', say the number in one sentence. "
        "You are also the official assistant for BotifyAI, a business customer-support chatbot service for websites and WhatsApp, and you know the following verified information about it. "
        "You are not the unrelated Botify AI character-roleplay app, and you must not describe its features, subscriptions, characters, or mobile app. "
        "If asked about that unrelated app, clearly say you are the business-support BotifyAI service and answer only about this service. "
        "BotifyAI business information: BotifyAI answers customer questions 24/7 on websites and WhatsApp, supports 22+ languages including Arabic, English, Hindi, Urdu, French, Spanish, Portuguese, and more. It can be trained on products, services, prices, FAQs, and policies, captures leads, and provides conversation and analytics insights. Works with WordPress, Shopify, Wix, Webflow, custom HTML, and most platforms that allow a small code snippet. Typical setup takes about 5 minutes; free installation help available by email at hello@botifyai.xyz. "
        "IMPORTANT: NEVER mention 'Indian languages', India-specific pricing, or Rupees UNLESS the visitor is on the India page (see region above) or the visitor explicitly asks about India. Do not name languages by region — say 'multilingual' or 'many languages' instead of 'Indian languages' unless India is the current region. "
        "All plans show a 7-day money-back guarantee. When talking specifically about BotifyAI's own plans, pricing, or features, do not promise anything beyond the region pricing shown above — if asked for account-specific help or something unknown about BotifyAI itself, collect the user's name, business, contact details, and question for the team instead of guessing. "
        "Talk like an intelligent, well-mannered human friend. "
        "You can speak many languages fluently. Whatever language the user writes or speaks in, reply directly in that same language. "
        "Switch languages instantly and naturally. Never say you can only speak one language. "
        "Never add English translations, never explain phrases, never explain emojis, and never quote the user's words back with definitions. "
        "Use common sense: a greeting like 'kya haal hai' means 'how are you' — simply reply naturally. "
        "When greeting the user or opening a reply with a greeting word, always use 'Hello' (never 'Namaste' or other language-specific greeting words), even when the rest of your reply is in Hindi or another language. "
        "When the user uses Hindi (or Roman Hindi), write your reply in Devanagari script (हिंदी). "
        "When the user uses Urdu, write your reply in Arabic/Persian script (اردو). "
        "When the user uses Arabic, write your reply in pure Arabic script (العربية الفصحى) only. "
        "When the user uses Bengali, write your reply in Bengali script. "
        "Stay short. Detailed only if the user asks. "
        "If you are unsure about something, say so honestly — in one line. "
        "AFTER answering an OFF-TOPIC question (anything not directly about this specific business the bot is deployed on), add ONE friendly line that gently pulls the conversation back to the business — like a curious short question about what they were looking for on this site, or a soft invitation. Keep this bridge under 15 words, warm, never salesy, never pushy. Skip the bridge entirely if the question WAS about the business — answering it is enough. Never bridge more than once every 3-4 turns; if you already bridged, just answer and stop. "
        f"The current date and time is: {now_str}. When asked about today's date, current time, or anything time-related, "
        "you must answer using this exact date and time. Do not say you lack real-time information. "
    )

def _detect_country_code():
    """Return an uppercase ISO country code for the current request, or None."""
    # 1. Cloudflare (if the site is behind CF)
    cc = request.headers.get("CF-IPCountry")
    if cc and cc.upper() not in ("XX", "T1"):
        return cc.upper()
    # 2. Vercel / other CDNs
    for h in ("X-Vercel-IP-Country", "X-Country-Code", "X-AppEngine-Country", "X-Country"):
        v = request.headers.get(h)
        if v:
            return v.upper()
    # 3. Fall back to IP lookup (Railway forwards real IP in X-Forwarded-For)
    try:
        fwd = request.headers.get("X-Forwarded-For", "")
        ip = fwd.split(",")[0].strip() if fwd else request.remote_addr
        if ip and ip not in ("127.0.0.1", "::1"):
            r = requests.get(f"https://ipapi.co/{ip}/country/", timeout=2)
            if r.ok:
                code = r.text.strip().upper()
                if len(code) == 2:
                    return code
    except Exception:
        pass
    return None

@app.route("/")
def home():
    # Skip auto-redirect if the visitor explicitly asked for the default page
    if request.args.get("noredirect") == "1":
        return app.send_static_file("india-landing.html")
    country = _detect_country_code()
    country_routes = {
        "AE": "/uae", "QA": "/qatar", "KW": "/kuwait", "BH": "/bahrain", "OM": "/oman",
        "BD": "/bangladesh", "LK": "/srilanka", "NP": "/nepal", "CN": "/china",
        "BT": "/bhutan", "MM": "/myanmar", "ID": "/indonesia",
        "PK": "/pakistan", "JO": "/jordan", "IQ": "/iraq", "LB": "/lebanon",
        "PH": "/philippines", "VN": "/vietnam", "MY": "/malaysia", "TH": "/thailand",
        "NG": "/nigeria", "KE": "/kenya", "ZA": "/southafrica", "EG": "/egypt",
        "GH": "/ghana", "UG": "/uganda", "TZ": "/tanzania", "RW": "/rwanda",
        "MX": "/mexico", "BR": "/brazil", "CO": "/colombia", "AR": "/argentina",
        "PE": "/peru", "CL": "/chile",
    }
    if country in country_routes:
        return redirect(country_routes[country], code=302)
    return app.send_static_file("india-landing.html")

@app.route("/widget.js")
def widget_js():
    return app.send_static_file("widget.js")

@app.route("/widget.html")
def widget_html():
    return app.send_static_file("widget.html")

@app.route("/india")
def india_landing():
    return app.send_static_file("india-landing.html")

@app.route("/uae")
def uae_landing():
    return app.send_static_file("uae-landing.html")

@app.route("/qatar")
def qatar_landing():
    return app.send_static_file("qatar-landing.html")

@app.route("/kuwait")
def kuwait_landing():
    return app.send_static_file("kuwait-landing.html")

@app.route("/bahrain")
def bahrain_landing():
    return app.send_static_file("bahrain-landing.html")

@app.route("/oman")
def oman_landing():
    return app.send_static_file("oman-landing.html")

@app.route("/bangladesh")
def bangladesh_landing():
    return app.send_static_file("bangladesh-landing.html")

@app.route("/srilanka")
def srilanka_landing():
    return app.send_static_file("srilanka-landing.html")

@app.route("/nepal")
def nepal_landing():
    return app.send_static_file("nepal-landing.html")

@app.route("/china")
def china_landing():
    return app.send_static_file("china-landing.html")

@app.route("/bhutan")
def bhutan_landing():
    return app.send_static_file("bhutan-landing.html")

@app.route("/myanmar")
def myanmar_landing():
    return app.send_static_file("myanmar-landing.html")

@app.route("/indonesia")
def indonesia_landing():
    return app.send_static_file("indonesia-landing.html")

def _make_landing(slug):
    def _view():
        return app.send_static_file(f"{slug}-landing.html")
    _view.__name__ = f"{slug}_landing"
    return _view

for _slug in [
    "pakistan","jordan","iraq","lebanon","philippines","vietnam","malaysia","thailand",
    "nigeria","kenya","southafrica","egypt","ghana","uganda","tanzania","rwanda",
    "mexico","brazil","colombia","argentina","peru","chile"
]:
    app.add_url_rule(f"/{_slug}", endpoint=f"{_slug}_landing", view_func=_make_landing(_slug))

@app.route("/payment")
def payment():
    return app.send_static_file("payment.html")

@app.route("/dashboard")
def dashboard():
    return app.send_static_file("dashboard.html")

@app.route("/terms")
def terms():
    return app.send_static_file("terms.html")

@app.route("/privacy")
def privacy():
    return app.send_static_file("privacy.html")

@app.route("/refund")
def refund():
    return app.send_static_file("refund.html")

@app.route("/admin")
def admin():
    return app.send_static_file("admin.html")

@app.route("/deploy-guide")
def deploy_guide():
    return app.send_static_file("deploy-guide.html")

@app.route("/india-deploy-guide")
def india_deploy_guide():
    return app.send_static_file("india-deploy-guide.html")

@app.route("/my-business")
def my_business():
    return app.send_static_file("my-business.html")

@app.route("/partners")
def partners():
    return app.send_static_file("partners.html")

@app.route("/onboard")
def onboard():
    return app.send_static_file("onboard.html")

@app.route("/widget-builder")
@app.route("/customize")
def widget_builder():
    return app.send_static_file("widget-builder.html")

@app.route("/my-bot")
@app.route("/login")
@app.route("/dashboard")
def my_bot_login():
    """Dedicated 'log in with access code' page. Verifies then redirects to /onboard?code=... for editing."""
    return app.send_static_file("my-bot.html")

@app.route("/api/upload-icon", methods=["POST"])
def upload_icon():
    """Client uploads a logo/photo → we host it and return a public URL.
    Avoids the data:URL problem where huge base64 blobs in a script-tag
    attribute get truncated by HTML parsers or sanitizers on the customer's site."""
    try:
        if 'file' not in request.files:
            return jsonify({"success": False, "message": "No file provided"}), 400
        f = request.files['file']
        if not f or not f.filename:
            return jsonify({"success": False, "message": "Empty file"}), 400
        # Size check — read raw to know length
        content = f.read()
        if len(content) > 800 * 1024:
            return jsonify({"success": False, "message": "Image is bigger than 800KB — please compress and try again."}), 400
        # Content-type sniff via extension only (small surface; we accept common web image formats)
        ext = os.path.splitext(f.filename.lower())[1]
        if ext not in {".png", ".jpg", ".jpeg", ".gif", ".webp", ".svg"}:
            return jsonify({"success": False, "message": "Only PNG, JPG, GIF, WEBP, or SVG allowed."}), 400
        upload_dir = os.path.join(os.path.dirname(__file__), "static", "uploads", "icons")
        os.makedirs(upload_dir, exist_ok=True)
        # Unique filename so uploads don't collide + never expose the original name
        safe_name = f"{uuid.uuid4().hex}{ext}"
        with open(os.path.join(upload_dir, safe_name), "wb") as out:
            out.write(content)
        public_url = f"/static/uploads/icons/{safe_name}"
        return jsonify({"success": True, "url": public_url})
    except Exception as e:
        logger.error(f"upload_icon error: {e}")
        return jsonify({"success": False, "message": "Upload failed. Please try again."}), 500

@app.route("/api/quick-access", methods=["POST"])
def quick_access():
    """Instant access code — client fills a small form, we create the account and
    hand back a code + onboarding link. No WhatsApp step, no admin approval."""
    try:
        data = request.get_json(silent=True) or {}
        name = (data.get("name") or "").strip()
        email = (data.get("email") or "").strip()
        website = (data.get("website") or "").strip()
        plan = (data.get("plan") or "starter").strip()
        region = (data.get("region") or "").strip().lower() or "uae"
        if not name or not email:
            return jsonify({"success": False, "message": "Name and email are required."}), 400
        # Normalise region to what create_client accepts
        if region not in {"india", "uae", "nri"}:
            region = "uae"
        client_id, access_code = create_client(name, email, website, plan, days_valid=30, region=region)
        # Fire notifications — welcome to the client, alert to the owner
        try:
            send_welcome_email(email, name, access_code, plan)
        except Exception as e:
            logger.error(f"quick_access welcome email failed: {e}")
        try:
            notify_owner(name, email, website, plan, access_code)
        except Exception as e:
            logger.error(f"quick_access notify_owner failed: {e}")
        return jsonify({
            "success": True,
            "access_code": access_code,
            "onboard_url": f"/onboard?code={access_code}",
        })
    except Exception as e:
        logger.error(f"quick_access error: {e}")
        return jsonify({"success": False, "message": "Server error, please try again."}), 500

@app.route("/onboard/verify", methods=["POST"])
def onboard_verify():
    """Check whether an access code is valid before showing the onboarding form."""
    try:
        data = request.get_json() or {}
        code = (data.get("access_code") or "").strip().upper()
        if not code:
            return jsonify({"valid": False, "message": "Please enter your access code."}), 400
        # Normalize: users often paste just the 16-hex suffix (from a shortened widget snippet,
        # a chat log, etc.) without the "BOT-" prefix — auto-prefix so lookup still works.
        if not code.startswith("BOT-") and len(code) == 16 and all(c in "0123456789ABCDEF" for c in code):
            code = "BOT-" + code
        client = get_client_by_access_code(code)
        if not client:
            return jsonify({"valid": False, "message": "This access code was not found."}), 404
        # Detect if already onboarded (any existing docs) and load the most recent saved payload
        # so the client-side form can pre-fill every field for editing/retraining.
        docs = get_documents(client["id"]) or []
        onboarded = any((d.get("filename") or "").startswith("onboard_") for d in docs)
        payload = None
        # get_documents returns newest first — grab the latest onboard_payload.json we find
        import json as _json
        for d in docs:
            if (d.get("filename") or "") == "onboard_payload.json":
                try:
                    payload = _json.loads(d.get("content") or "{}")
                    break
                except Exception:
                    continue
        return jsonify({
            "valid": True,
            "name": client.get("name", ""),
            "website": client.get("website", ""),
            "already_onboarded": onboarded,
            "payload": payload,  # null if fresh signup; full form data if returning to edit
        })
    except Exception as e:
        logger.error(f"onboard/verify error: {e}")
        return jsonify({"valid": False, "message": "Server error, please try again."}), 500

@app.route("/onboard/submit", methods=["POST"])
def onboard_submit():
    """Save the onboarding form's answers as bot training documents."""
    try:
        raw = request.form.get("data")
        if not raw:
            return jsonify({"success": False, "message": "Missing data"}), 400
        import json as _json
        payload = _json.loads(raw)
        code = (payload.get("access_code") or "").strip().upper()
        if not code:
            return jsonify({"success": False, "message": "Missing access code"}), 400
        client = get_client_by_access_code(code)
        if not client:
            return jsonify({"success": False, "message": "Access code not found"}), 404
        client_id = client["id"]

        # Compose a single training document from all fields
        parts = [f"BUSINESS: {payload.get('name','')}"]
        if payload.get("business_type"): parts.append(f"Type: {payload['business_type']}")
        if payload.get("city"): parts.append(f"City: {payload['city']}")
        if payload.get("address"): parts.append(f"Address: {payload['address']}")
        if payload.get("phone"): parts.append(f"Phone: {payload['phone']}")
        if payload.get("hours"): parts.append(f"Hours: {payload['hours']}")
        if payload.get("website"): parts.append(f"Website: {payload['website']}")
        if payload.get("description"): parts.append(f"\nABOUT:\n{payload['description']}")
        if payload.get("products"): parts.append(f"\nPRODUCTS / SERVICES / MENU:\n{payload['products']}")
        if payload.get("faqs"):
            parts.append("\nFREQUENTLY ASKED QUESTIONS:")
            for i, f in enumerate(payload["faqs"], 1):
                parts.append(f"\nQ{i}: {f.get('q','')}\nA{i}: {f.get('a','')}")
        if payload.get("refund_policy"): parts.append(f"\nREFUND / CANCELLATION POLICY:\n{payload['refund_policy']}")
        if payload.get("delivery_info"): parts.append(f"\nDELIVERY / BOOKING INFO:\n{payload['delivery_info']}")
        if payload.get("languages"): parts.append(f"\nLANGUAGES BOT SHOULD REPLY IN: {', '.join(payload['languages'])}")
        if payload.get("extra_notes"): parts.append(f"\nADDITIONAL NOTES:\n{payload['extra_notes']}")
        text_content = "\n".join(parts)

        # Save two docs in parallel:
        #   1. onboard_business_info.txt — human-readable free text used to train the bot
        #   2. onboard_payload.json — the original structured payload so we can reload every
        #      field verbatim when the client comes back to edit their profile
        try:
            add_document(client_id, "onboard_business_info.txt", text_content, 0)
        except Exception as doc_err:
            logger.error(f"onboard/submit add_document failed: {doc_err}")
        try:
            add_document(client_id, "onboard_payload.json", _json.dumps(payload), 0)
        except Exception as doc_err:
            logger.error(f"onboard/submit payload save failed: {doc_err}")

        # Save any uploaded files as additional documents
        for key in request.files:
            f = request.files[key]
            if not f or not f.filename:
                continue
            try:
                content = f.read()
                # For text-like files store as text; PDF/DOC left to admin extraction later
                if f.filename.lower().endswith((".txt",)):
                    add_document(client_id, f"onboard_{f.filename}", content.decode("utf-8", errors="ignore"), 0)
                else:
                    # Save the file blob to disk for admin to process
                    upload_dir = os.path.join(os.path.dirname(__file__), "uploads")
                    os.makedirs(upload_dir, exist_ok=True)
                    safe_name = f"onboard_{client_id}_{f.filename}".replace("..", "_")
                    path = os.path.join(upload_dir, safe_name)
                    with open(path, "wb") as out:
                        out.write(content)
                    add_document(client_id, f"onboard_{f.filename}", f"[Uploaded file — admin will extract] path={safe_name}", 0)
            except Exception as inner:
                logger.error(f"onboard file save failed: {inner}")

        # Notify owner that onboarding is complete
        try:
            notify_onboard_complete(client, payload)
        except Exception as ne:
            logger.error(f"onboard notify failed: {ne}")

        return jsonify({"success": True, "message": "Bot training data saved"})
    except Exception as e:
        logger.error(f"onboard/submit error: {e}")
        return jsonify({"success": False, "message": "Server error — please email hello@botifyai.xyz"}), 500


def notify_onboard_complete(client, payload):
    """Send a WhatsApp + email nudge to the owner that a client finished onboarding."""
    name = client.get("name") or payload.get("name", "")
    plan = client.get("plan", "")
    biz_type = payload.get("business_type", "")
    city = payload.get("city", "")
    owner_phone = os.environ.get("OWNER_WHATSAPP", "")
    twilio_sid = os.environ.get("TWILIO_ACCOUNT_SID", "")
    twilio_tok = os.environ.get("TWILIO_AUTH_TOKEN", "")
    twilio_wa = os.environ.get("TWILIO_WHATSAPP_NUMBER", "whatsapp:+14155238886")
    if owner_phone and twilio_sid and twilio_tok:
        try:
            from twilio.rest import Client as TwilioClient
            tc = TwilioClient(twilio_sid, twilio_tok)
            tc.messages.create(
                body=(
                    f"✅ ONBOARDING COMPLETE\n\n"
                    f"👤 {name}\n"
                    f"🏢 {biz_type} · {city}\n"
                    f"📦 {plan.title()}\n"
                    f"🔑 {client.get('access_code','')}\n\n"
                    f"Review docs: /admin"
                ),
                from_=twilio_wa,
                to=f"whatsapp:{owner_phone}",
            )
        except Exception as e:
            logger.error(f"onboard notify whatsapp failed: {e}")



def check_admin_password(data_or_args):
    pw = (data_or_args.get("pw", "") if isinstance(data_or_args, dict) else data_or_args.get("pw", ""))
    return pw == os.environ.get("ADMIN_PASSWORD", "farman2024")

@app.route("/admin/clients", methods=["GET"])
def admin_clients():
    if not check_admin_password(request.args):
        return jsonify({"error": "Unauthorized"}), 401
    return jsonify({"clients": get_all_clients()})

@app.route("/admin/clients", methods=["POST"])
def admin_create_client():
    data = request.get_json(silent=True) or {}
    if not check_admin_password(data):
        return jsonify({"error": "Unauthorized"}), 401
    name = data.get("name", "").strip()
    email = data.get("email", "").strip()
    website = data.get("website", "").strip()
    plan = data.get("plan", "basic").strip()
    region = data.get("region", "uae").strip().lower()
    if region not in {"india", "uae"}:
        region = "uae"
    days = int(data.get("days_valid", 365) or 365)
    if not name:
        return jsonify({"error": "Name is required"}), 400
    client_id, access_code = create_client(name, email, website, plan, days, region)
    # Send welcome email to the new client and notify the owner — same flow as automated payment
    try:
        if email:
            send_welcome_email(email, name, access_code, plan)
    except Exception as e:
        logger.error(f"admin_create_client: welcome email failed: {e}")
    try:
        notify_owner(name, email, website, plan, access_code)
    except Exception as e:
        logger.error(f"admin_create_client: notify_owner failed: {e}")
    return jsonify({"success": True, "client_id": client_id, "access_code": access_code})

@app.route("/admin/clients/<int:client_id>", methods=["DELETE"])
def admin_delete_client(client_id):
    pw = request.args.get("pw", "")
    if pw != os.environ.get("ADMIN_PASSWORD", "farman2024"):
        return jsonify({"error": "Unauthorized"}), 401
    from database import delete_client
    delete_client(client_id)
    return jsonify({"success": True})

@app.route("/admin/clients/<int:client_id>/documents", methods=["GET"])
def admin_client_documents(client_id):
    if not check_admin_password(request.args):
        return jsonify({"error": "Unauthorized"}), 401
    return jsonify({"documents": get_documents(client_id)})

@app.route("/admin/clients/<int:client_id>/documents", methods=["POST"])
def admin_upload_client_document(client_id):
    pw = request.form.get("pw", "")
    if pw != os.environ.get("ADMIN_PASSWORD", "farman2024"):
        return jsonify({"error": "Unauthorized"}), 401
    if 'file' not in request.files:
        return jsonify({"error": "No file provided"}), 400
    file = request.files['file']
    if file.filename == '':
        return jsonify({"error": "No file selected"}), 400
    client = get_client(client_id)
    if not client:
        return jsonify({"error": "Client not found"}), 404
    file.seek(0, os.SEEK_END)
    file_size = file.tell()
    file.seek(0)
    logger.info(f"Upload: client={client_id} file={file.filename} size={file_size} bytes")
    if file_size > 5 * 1024 * 1024:
        return jsonify({"error": "File too large. Maximum size is 5MB. Please upload a smaller text-based PDF, DOCX, or TXT file."}), 413
    try:
        text = extract_text_from_file(file, file.filename)
    except Exception as e:
        logger.error("Extraction failed: %s", e)
        return jsonify({"error": "Failed to extract text from file: " + str(e)}), 500
    text_len = len(text.strip())
    logger.info(f"Extracted text length: {text_len} chars")
    if not text.strip():
        return jsonify({"error": "No text could be extracted. If this is a scanned/image PDF, please upload a text-based PDF, DOCX, or TXT file."}), 400
    chunks = chunk_text(text)
    logger.info(f"Created {len(chunks)} chunks")
    doc_id = add_document(client_id, file.filename, text, len(chunks))
    return jsonify({"success": True, "doc_id": doc_id, "chunks": len(chunks), "preview": text[:300].replace('\n', ' ')})

@app.route("/admin/clients/<int:client_id>/documents/text", methods=["POST"])
def admin_add_text_document(client_id):
    """Add a document directly from pasted text."""
    data = request.get_json() or request.form
    pw = data.get("pw", "") if request.get_json() else request.form.get("pw", "")
    if pw != os.environ.get("ADMIN_PASSWORD", "farman2024"):
        return jsonify({"error": "Unauthorized"}), 401
    client = get_client(client_id)
    if not client:
        return jsonify({"error": "Client not found"}), 404
    title = (data.get("title") or "Pasted text").strip()
    content = (data.get("content") or "").strip()
    if not content:
        return jsonify({"error": "Content is required"}), 400
    if len(content) > 100000:
        return jsonify({"error": "Content too large. Max 100,000 characters."}), 413
    chunks = chunk_text(content)
    logger.info(f"Text document: client={client_id} title={title} chars={len(content)} chunks={len(chunks)}")
    doc_id = add_document(client_id, title + ".txt", content, len(chunks))
    return jsonify({"success": True, "doc_id": doc_id, "chunks": len(chunks), "preview": content[:300].replace('\n', ' ')})

@app.route("/save-lead", methods=["POST"])
def save_lead():
    try:
        data = request.get_json()
        
        # Validate required fields
        required_fields = ['name', 'email', 'phone']
        for field in required_fields:
            if not data.get(field, '').strip():
                return jsonify({"error": f"Missing required field: {field}"}), 400
        
        # Create lead record
        lead = {
            "id": len(LEADS) + 1,
            "name": data['name'].strip(),
            "email": data['email'].strip().lower(),
            "phone": data['phone'].strip(),
            "company": data.get('company', '').strip(),
            "timestamp": data.get('timestamp', datetime.now().isoformat()),
            "source": data.get('source', 'unknown'),
            "status": "new"
        }
        
        # Check for duplicate leads
        for existing_lead in LEADS:
            if (existing_lead['email'] == lead['email'] or 
                existing_lead['phone'] == lead['phone']):
                # Update existing lead instead of creating duplicate
                existing_lead.update(lead)
                logger.info(f"Updated existing lead: {lead['email']}")
                return jsonify({"success": True, "lead_id": existing_lead['id'], "updated": True})
        
        # Add new lead
        LEADS.append(lead)
        logger.info(f"New lead captured: {lead['name']} - {lead['email']}")
        
        return jsonify({
            "success": True, 
            "lead_id": lead['id'],
            "message": "Lead saved successfully"
        })
        
    except Exception as e:
        logger.error(f"Error saving lead: {e}")
        return jsonify({"error": "Failed to save lead"}), 500

@app.route("/leads", methods=["GET"])
def get_leads():
    # Simple endpoint to view leads (in production, add authentication)
    return jsonify({
        "leads": LEADS,
        "total": len(LEADS)
    })

@app.route("/create-payment", methods=["POST"])
def create_payment():
    """Create Razorpay payment order for Indian customers"""
    try:
        data = request.get_json()
        plan = data.get('plan', 'starter')
        
        # Plan pricing in INR
        pricing = {
            'starter': {'amount': 10000, 'name': 'Starter Plan'},  # ₹100 in paise
            'professional': {'amount': 15000, 'name': 'Professional Plan'},  # ₹150 in paise
            'enterprise': {'amount': 20000, 'name': 'Enterprise Plan'}  # ₹200 in paise
        }
        
        if plan not in pricing:
            return jsonify({"error": "Invalid plan selected"}), 400
        
        plan_details = pricing[plan]
        
        # In production, integrate with actual Razorpay API
        # For now, return mock payment order
        payment_order = {
            "id": f"order_{datetime.now().strftime('%Y%m%d_%H%M%S')}",
            "amount": plan_details['amount'],
            "currency": "INR",
            "name": "AI Chatbot India",
            "description": plan_details['name'],
            "customer_email": data.get('email', ''),
            "customer_phone": data.get('phone', ''),
            "callback_url": f"{request.url_root}payment-success",
            "notes": {
                "plan": plan,
                "customer_name": data.get('name', '')
            }
        }
        
        logger.info(f"Payment order created: {payment_order['id']} for plan {plan}")
        return jsonify({
            "success": True,
            "order": payment_order,
            "key": "rzp_test_1234567890"  # Test key in production
        })
        
    except Exception as e:
        logger.error(f"Error creating payment: {e}")
        return jsonify({"error": "Failed to create payment"}), 500

@app.route("/payment-success", methods=["POST"])
def payment_success():
    """Handle successful payment callback and grant automatic access"""
    try:
        data = request.get_json()
        payment_id = data.get('payment_id')
        order_id = data.get('order_id')
        plan = data.get('plan', 'starter')
        email = data.get('email', '')
        name = data.get('name', '')
        website = data.get('website', '')
        access_code_in = data.get('access_code', '')

        logger.info(f"Payment successful: {payment_id} for order {order_id} - Plan: {plan}")

        # Grant automatic access — create client in database
        client_id, access_code = create_client(name, email, website, plan, days_valid=30)
        
        # In production, save to database
        logger.info(f"Access granted: {access_code} for {email}")
        
        # Send access details via email (in production)
        send_access_email(email, name, access_code, plan)
        
        client = get_client(client_id)
        return jsonify({
            "success": True,
            "message": "Payment processed successfully",
            "access_code": access_code,
            "redirect_url": f"/chat?access={access_code}",
            "plan": plan,
            "expires_at": client['expires_at']
        })
        
    except Exception as e:
        logger.error(f"Error processing payment success: {e}")
        return jsonify({"error": "Payment processing failed"}), 500

def send_access_email(email, name, access_code, plan):
    """Send access details via email (in production)"""
    logger.info(f"Email sent to {email}: Access Code: {access_code}, Plan: {plan}")
    send_welcome_email(email, name, access_code, plan)


def send_welcome_email(to_email, name, access_code, plan):
    """Send beautiful welcome email with embed code to new client"""
    gmail_user = os.environ.get('GMAIL_USER', '')
    gmail_pass = os.environ.get('GMAIL_APP_PASSWORD', '')
    if not gmail_user or not gmail_pass:
        logger.warning('GMAIL_USER or GMAIL_APP_PASSWORD not set — skipping welcome email')
        return

    plan_label  = plan.title() if plan else 'Paid'
    price       = ''  # currency depends on region — bot mentions it after onboarding instead
    site_origin = os.environ.get('PUBLIC_BASE_URL', 'https://botifyai.xyz')
    embed_code  = f'<script src="{site_origin}/widget.js" data-agent="AI Assistant" data-welcome="Hello! How can I help you?" data-access="{access_code}"></script>'
    dashboard_url = f'{site_origin}/onboard?code={access_code}'
    guide_url = f'{site_origin}/deploy-guide?access={access_code}'

    html = f"""
    <!DOCTYPE html>
    <html>
    <head><meta charset="UTF-8"><meta name="viewport" content="width=device-width,initial-scale=1"></head>
    <body style="margin:0;padding:0;background:#0a0a1a;font-family:Inter,Arial,sans-serif;">
      <div style="max-width:600px;margin:0 auto;padding:40px 20px;">

        <div style="text-align:center;margin-bottom:32px;">
          <div style="font-size:48px;">🤖</div>
          <h1 style="color:#a78bfa;font-size:1.8rem;margin:12px 0 4px;">Your AI Bot is Ready!</h1>
          <p style="color:#64748b;margin:0;">Payment confirmed — here's everything you need</p>
        </div>

        <div style="background:rgba(255,255,255,.05);border:1px solid rgba(167,139,250,.3);border-radius:16px;padding:24px;margin-bottom:20px;">
          <p style="color:#94a3b8;margin:0 0 4px;font-size:.85rem;text-transform:uppercase;letter-spacing:.5px;">Welcome</p>
          <p style="color:#e2e8f0;font-size:1.1rem;font-weight:600;margin:0;">Hi {name}!</p>
        </div>

        <div style="background:rgba(255,255,255,.05);border:1px solid rgba(167,139,250,.3);border-radius:16px;padding:24px;margin-bottom:20px;">
          <p style="color:#94a3b8;margin:0 0 8px;font-size:.85rem;text-transform:uppercase;letter-spacing:.5px;">Your Access Code</p>
          <div style="background:#0d0d1f;border-radius:8px;padding:14px;font-family:monospace;font-size:1.3rem;color:#a78bfa;letter-spacing:2px;text-align:center;">{access_code}</div>
          <p style="color:#475569;font-size:.8rem;margin:8px 0 0;text-align:center;">Keep this safe — it's your key to the dashboard</p>
        </div>

        <div style="background:rgba(255,255,255,.05);border:1px solid rgba(167,139,250,.3);border-radius:16px;padding:24px;margin-bottom:20px;">
          <p style="color:#94a3b8;margin:0 0 8px;font-size:.85rem;text-transform:uppercase;letter-spacing:.5px;">Plan</p>
          <p style="color:#34d399;font-size:1.1rem;font-weight:700;margin:0;">{plan_label} Plan · Activated</p>
        </div>

        <div style="background:linear-gradient(135deg,rgba(167,139,250,.15),rgba(96,165,250,.1));border:1px solid rgba(167,139,250,.35);border-radius:16px;padding:24px;margin-bottom:20px;text-align:center;">
          <p style="color:#e2e8f0;margin:0 0 12px;font-size:.95rem;font-weight:600;">🚀 Finish setup in 5 minutes</p>
          <p style="color:#94a3b8;margin:0 0 14px;font-size:.85rem;line-height:1.6;">Answer a few quick questions about your business — your bot learns and goes live automatically.</p>
          <a href="{dashboard_url}" style="display:inline-block;background:linear-gradient(135deg,#a78bfa,#60a5fa);color:#fff;text-decoration:none;padding:12px 28px;border-radius:9px;font-weight:700;font-size:.95rem;">Complete Setup →</a>
        </div>

        <div style="background:rgba(255,255,255,.05);border:1px solid rgba(167,139,250,.3);border-radius:16px;padding:24px;margin-bottom:20px;">
          <p style="color:#94a3b8;margin:0 0 12px;font-size:.85rem;text-transform:uppercase;letter-spacing:.5px;">🚀 Add Bot to Your Website</p>
          <p style="color:#94a3b8;font-size:.9rem;margin:0 0 12px;">Copy this one line and paste it in your website HTML before <code style="color:#a78bfa;">&lt;/body&gt;</code>:</p>
          <div style="background:#0d0d1f;border-radius:8px;padding:14px;font-family:monospace;font-size:.78rem;color:#60a5fa;word-break:break-all;line-height:1.6;">{embed_code}</div>
        </div>

        <div style="background:rgba(255,255,255,.05);border:1px solid rgba(167,139,250,.3);border-radius:16px;padding:24px;margin-bottom:28px;text-align:center;">
          <p style="color:#94a3b8;margin:0 0 6px;font-size:.85rem;">Step-by-step guide for WordPress, Shopify, Wix &amp; more:</p>
          <a href="{guide_url}" style="display:inline-block;background:rgba(167,139,250,.15);border:1px solid rgba(167,139,250,.3);color:#a78bfa;text-decoration:none;padding:12px 28px;border-radius:9px;font-weight:700;font-size:.95rem;">📖 View Full Deploy Guide →</a>
          <p style="color:#475569;font-size:.8rem;margin:12px 0 0;">Can't do it? Reply to this email — we'll add it for you free! ✅</p>
        </div>

        <div style="text-align:center;margin-bottom:32px;display:flex;gap:12px;justify-content:center;flex-wrap:wrap;">
          <a href="{dashboard_url}" style="display:inline-block;background:linear-gradient(135deg,#a78bfa,#60a5fa);color:#fff;text-decoration:none;padding:15px 36px;border-radius:10px;font-weight:700;font-size:1rem;">Go to My Dashboard →</a>
        </div>

        <p style="color:#334155;font-size:.82rem;text-align:center;">Questions? Reply to this email or message us on <a href="https://wa.me/917992111021" style="color:#60a5fa;">WhatsApp</a></p>
      </div>
    </body>
    </html>
    """

    reply_to = os.environ.get('SUPPORT_EMAIL', 'support@botifyai.xyz')
    try:
        msg = MIMEMultipart('alternative')
        msg['Subject']  = f'🤖 Your BotifyAI Bot is Ready — {plan_label} Plan Activated!'
        msg['From']     = f'BotifyAI Support <{gmail_user}>'
        msg['Reply-To'] = f'BotifyAI Support <{reply_to}>'
        msg['To']       = to_email
        msg.attach(MIMEText(html, 'html'))

        with smtplib.SMTP_SSL('smtp.gmail.com', 465) as server:
            server.login(gmail_user, gmail_pass)
            server.sendmail(gmail_user, to_email, msg.as_string())
        logger.info(f'Welcome email sent to {to_email}')
    except Exception as e:
        logger.error(f'Failed to send welcome email to {to_email}: {e}')

def notify_owner(name, email, website, plan, access_code):
    """Notify owner via WhatsApp + email when a new client pays"""
    plan_label = (plan or '').title() or 'Paid'
    price = ''  # currency inferred from region — kept blank here so old AED figures don't leak
    owner_phone = os.environ.get('OWNER_WHATSAPP', '')
    twilio_sid  = os.environ.get('TWILIO_ACCOUNT_SID', '')
    twilio_tok  = os.environ.get('TWILIO_AUTH_TOKEN', '')
    twilio_wa   = os.environ.get('TWILIO_WHATSAPP_NUMBER', 'whatsapp:+14155238886')

    # WhatsApp notification
    if owner_phone and twilio_sid and twilio_tok:
        try:
            from twilio.rest import Client as TwilioClient
            tc = TwilioClient(twilio_sid, twilio_tok)
            tc.messages.create(
                body=(
                    f"🎉 NEW CLIENT PAID!\n\n"
                    f"👤 Name: {name}\n"
                    f"📧 Email: {email}\n"
                    f"🌐 Website: {website or 'Not provided'}\n"
                    f"📦 Plan: {plan_label}\n"
                    f"🔑 Code: {access_code}\n\n"
                    f"👉 Admin: {os.environ.get('PUBLIC_BASE_URL','https://botifyai.xyz')}/admin"
                ),
                from_=twilio_wa,
                to=f'whatsapp:{owner_phone}'
            )
            logger.info(f'Owner WhatsApp notification sent for {email}')
        except Exception as e:
            logger.error(f'Owner WhatsApp notify failed: {e}')

    # Email notification to owner
    gmail_user = os.environ.get('GMAIL_USER', '')
    gmail_pass = os.environ.get('GMAIL_APP_PASSWORD', '')
    owner_email = os.environ.get('OWNER_EMAIL', gmail_user)
    if gmail_user and gmail_pass and owner_email:
        try:
            msg = MIMEMultipart('alternative')
            msg['Subject'] = f'🎉 New Client — {plan_label} — {name}'
            msg['From']    = f'BotifyAI Admin <{gmail_user}>'
            msg['To']      = owner_email
            body = f"""
            <div style="font-family:Arial;padding:20px;background:#0a0a1a;color:#e2e8f0;">
            <h2 style="color:#34d399;">🎉 New Client Registered!</h2>
            <table style="border-collapse:collapse;width:100%;">
              <tr><td style="padding:8px;color:#94a3b8;">Name</td><td style="padding:8px;color:#fff;font-weight:600;">{name}</td></tr>
              <tr><td style="padding:8px;color:#94a3b8;">Email</td><td style="padding:8px;color:#60a5fa;">{email}</td></tr>
              <tr><td style="padding:8px;color:#94a3b8;">Website</td><td style="padding:8px;color:#60a5fa;">{website or '—'}</td></tr>
              <tr><td style="padding:8px;color:#94a3b8;">Plan</td><td style="padding:8px;color:#a78bfa;font-weight:700;">{plan_label}</td></tr>
              <tr><td style="padding:8px;color:#94a3b8;">Access Code</td><td style="padding:8px;font-family:monospace;color:#a78bfa;">{access_code}</td></tr>
            </table>
            <br>
            <a href="{os.environ.get('PUBLIC_BASE_URL','https://botifyai.xyz')}/admin" style="background:linear-gradient(135deg,#a78bfa,#60a5fa);color:#fff;padding:12px 24px;border-radius:8px;text-decoration:none;font-weight:700;">Open Admin Panel →</a>
            </div>
            """
            msg.attach(MIMEText(body, 'html'))
            with smtplib.SMTP_SSL('smtp.gmail.com', 465) as server:
                server.login(gmail_user, gmail_pass)
                server.sendmail(gmail_user, owner_email, msg.as_string())
            logger.info(f'Owner email notification sent for {email}')
        except Exception as e:
            logger.error(f'Owner email notify failed: {e}')


@app.route("/verify-access", methods=["POST"])
def verify_access():
    """Verify access code and grant chatbot access"""
    try:
        data = request.get_json()
        access_code = data.get('access_code')
        
        if not access_code:
            return jsonify({"error": "Access code required"}), 400
        
        client = get_client_by_access_code(access_code)
        if client:
            return jsonify({
                "success": True,
                "message": "Access granted",
                "plan": client['plan'],
                "expires_at": client['expires_at']
            })
        else:
            return jsonify({"error": "Invalid access code"}), 401
            
    except Exception as e:
        logger.error(f"Error verifying access: {e}")
        return jsonify({"error": "Access verification failed"}), 500

@app.route("/payment-instructions", methods=["GET"])
def payment_instructions():
    """Display payment instructions with account details"""
    return jsonify({
        "success": True,
        "payment_methods": {
            "stripe": {
                "enabled": true,
                "publishable_key": "pk_live_YOUR_STRIPE_KEY",
                "supported_cards": ["Visa", "Mastercard", "RuPay", "Amex"],
                "currencies": ["INR", "USD", "EUR", "GBP", "AED", "SGD"]
            },
            "paypal": {
                "enabled": true,
                "client_id": "YOUR_PAYPAL_CLIENT_ID",
                "supported_countries": 195
            },
            "upi": {
                "enabled": true,
                "apps": ["PhonePe", "Google Pay", "Paytm", "BHIM"]
            }
        },
        "instructions": {
            "step1": "Choose your plan (Starter: ₹3,000, Professional: ₹10,000, Enterprise: ₹25,000)",
            "step2": "Make payment to any of the above methods",
            "step3": "Send payment screenshot with your email to payments@aichatbot.in",
            "step4": "Receive access code within 2 hours",
            "step5": "Use access code to activate your AI chatbot"
        },
        "note": "Please include your email address in payment description for faster processing"
    })

@app.route("/upi-payment", methods=["POST"])
def upi_payment():
    """Handle UPI payment requests"""
    try:
        data = request.get_json()
        plan = data.get('plan', 'starter')
        
        pricing = {
            'starter': 3000,
            'professional': 10000,
            'enterprise': 25000
        }
        
        if plan not in pricing:
            return jsonify({"error": "Invalid plan selected"}), 400
        
        amount = pricing[plan]
        
        # UPI details for payment
        upi_details = {
            "upi_id": "aichatbot@okicici",
            "amount": amount,
            "note": f"AI Chatbot {plan.title()} Plan",
            "customer_name": data.get('name', ''),
            "customer_phone": data.get('phone', ''),
            "transaction_id": f"UPI_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
        }
        
        logger.info(f"UPI payment initiated: {upi_details['transaction_id']}")
        
        return jsonify({
            "success": True,
            "upi_details": upi_details,
            "qr_code_url": f"https://api.qrserver.com/v1/create-qr-code/?size=200x200&data=upi://pay?pa={upi_details['upi_id']}&pn=AI%20Chatbot&am={amount}&cu=INR&tn={upi_details['note']}"
        })
        
    except Exception as e:
        logger.error(f"Error creating UPI payment: {e}")
        return jsonify({"error": "Failed to create UPI payment"}), 500

@app.route("/wise-payment", methods=["POST"])
def wise_payment():
    """Handle Wise (TransferWise) international payment requests"""
    try:
        data = request.get_json()
        plan = data.get('plan', 'starter')
        country = data.get('country', 'IN')
        
        # International pricing in local currencies
        pricing = {
            'IN': {'starter': 3000, 'professional': 10000, 'enterprise': 25000, 'currency': 'INR'},
            'AE': {'starter': 499, 'professional': 1499, 'enterprise': 3499, 'currency': 'AED'},
            'SG': {'starter': 89, 'professional': 269, 'enterprise': 629, 'currency': 'SGD'},
            'GB': {'starter': 49, 'professional': 149, 'enterprise': 349, 'currency': 'GBP'},
            'US': {'starter': 59, 'professional': 179, 'enterprise': 419, 'currency': 'USD'}
        }
        
        if country not in pricing:
            return jsonify({"error": "Country not supported"}), 400
        
        if plan not in ['starter', 'professional', 'enterprise']:
            return jsonify({"error": "Invalid plan selected"}), 400
        
        country_pricing = pricing[country]
        amount = country_pricing[plan]
        currency = country_pricing['currency']
        
        # Wise payment details
        wise_details = {
            "amount": amount,
            "currency": currency,
            "recipient_name": "AI Chatbot Services",
            "recipient_email": "payments@aichatbot.in",
            "wise_account_id": "AI123456789",
            "reference": f"Chatbot_{plan}_{datetime.now().strftime('%Y%m%d_%H%M%S')}",
            "customer_name": data.get('name', ''),
            "customer_email": data.get('email', ''),
            "description": f"AI Chatbot {plan.title()} Plan - {country}",
            "payment_url": f"https://wise.com/pay/{wise_details['reference']}",
            "bank_details": {
                "account_name": "AI Chatbot Services",
                "account_number": "1234567890",
                "ifsc_code": "ICIC0001234",
                "bank_name": "ICICI Bank",
                "branch": "Mumbai Main"
            }
        }
        
        logger.info(f"Wise payment initiated: {wise_details['reference']} for {amount} {currency}")
        
        return jsonify({
            "success": True,
            "payment_method": "wise",
            "details": wise_details,
            "instructions": f"Transfer {amount} {currency} to the provided bank details or use the Wise payment link"
        })
        
    except Exception as e:
        logger.error(f"Error creating Wise payment: {e}")
        return jsonify({"error": "Failed to create Wise payment"}), 500

@app.route("/paypal-payment", methods=["POST"])
def paypal_payment():
    """Handle PayPal international payment requests"""
    try:
        data = request.get_json()
        plan = data.get('plan', 'starter')
        country = data.get('country', 'IN')
        
        # PayPal pricing in local currencies
        pricing = {
            'IN': {'starter': 3000, 'professional': 10000, 'enterprise': 25000, 'currency': 'INR'},
            'AE': {'starter': 499, 'professional': 1499, 'enterprise': 3499, 'currency': 'AED'},
            'SG': {'starter': 89, 'professional': 269, 'enterprise': 629, 'currency': 'SGD'},
            'GB': {'starter': 49, 'professional': 149, 'enterprise': 349, 'currency': 'GBP'},
            'US': {'starter': 59, 'professional': 179, 'enterprise': 419, 'currency': 'USD'}
        }
        
        if country not in pricing:
            return jsonify({"error": "Country not supported"}), 400
        
        if plan not in ['starter', 'professional', 'enterprise']:
            return jsonify({"error": "Invalid plan selected"}), 400
        
        country_pricing = pricing[country]
        amount = country_pricing[plan]
        currency = country_pricing['currency']
        
        # PayPal payment details
        paypal_details = {
            "amount": amount,
            "currency": currency,
            "merchant_id": "AI_CHATBOT_MERCHANT",
            "paypal_email": "payments@aichatbot.in",
            "item_name": f"AI Chatbot {plan.title()} Plan",
            "item_number": f"CHATBOT_{plan.upper()}_{datetime.now().strftime('%Y%m%d_%H%M%S')}",
            "custom": f"customer:{data.get('email', '')};plan:{plan};country:{country}",
            "return_url": f"{request.url_root}payment-success",
            "cancel_url": f"{request.url_root}payment-cancelled",
            "notify_url": f"{request.url_root}paypal-ipn",
            "payment_url": f"https://www.paypal.com/cgi-bin/webscr?cmd=_xclick&business={paypal_details['paypal_email']}&item_name={paypal_details['item_name']}&amount={amount}&currency_code={currency}&return={paypal_details['return_url']}&cancel_return={paypal_details['cancel_url']}"
        }
        
        logger.info(f"PayPal payment initiated: {paypal_details['item_number']} for {amount} {currency}")
        
        return jsonify({
            "success": True,
            "payment_method": "paypal",
            "details": paypal_details,
            "redirect_url": paypal_details['payment_url']
        })
        
    except Exception as e:
        logger.error(f"Error creating PayPal payment: {e}")
        return jsonify({"error": "Failed to create PayPal payment"}), 500

@app.route("/razorpay-order", methods=["POST"])
def razorpay_order():
    try:
        import razorpay
        data = request.get_json()
        plan  = data.get('plan', 'starter')
        name  = data.get('name', '')
        email = data.get('email', '')
        # India pricing in paise (₹100, ₹150, ₹200)
        prices = {'starter': 10000, 'professional': 15000, 'enterprise': 20000}
        amount = prices.get(plan, 10000)
        client = razorpay.Client(auth=(
            os.environ.get("RAZORPAY_KEY_ID", ""),
            os.environ.get("RAZORPAY_KEY_SECRET", "")
        ))
        order = client.order.create({
            'amount': amount,
            'currency': 'INR',
            'notes': {'plan': plan, 'name': name, 'email': email}
        })
        return jsonify({
            'success': True,
            'order_id': order['id'],
            'amount': amount,
            'currency': 'INR',
            'key_id': os.environ.get("RAZORPAY_KEY_ID", "")
        })
    except Exception as e:
        logger.error(f"Razorpay order error: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route("/stripe-checkout", methods=["POST"])
def stripe_checkout():
    """Create Stripe Checkout session and return URL"""
    try:
        data = request.get_json()
        plan     = data.get('plan', 'starter')
        name     = data.get('name', '')
        email    = data.get('email', '')
        currency = data.get('currency', 'aed')

        # AED pricing (in fils = AED * 100)
        aed_prices = {'starter': 49900, 'professional': 149900, 'enterprise': 349900}
        inr_prices = {'starter': 300000, 'professional': 1000000, 'enterprise': 2500000}

        stripe_key = os.environ.get("STRIPE_SECRET_KEY", "")

        if stripe_key:
            import stripe as stripe_lib
            stripe_lib.api_key = stripe_key
            amount   = aed_prices.get(plan, 49900) if currency == 'aed' else inr_prices.get(plan, 300000)
            cur_code = 'aed' if currency == 'aed' else 'inr'
            session  = stripe_lib.checkout.Session.create(
                payment_method_types=['card'],
                line_items=[{
                    'price_data': {
                        'currency': cur_code,
                        'unit_amount': amount,
                        'product_data': {
                            'name': f'AI Chatbot {plan.title()} Plan',
                            'description': f'Monthly subscription — {plan.title()} plan'
                        },
                    },
                    'quantity': 1,
                }],
                mode='payment',
                customer_email=email,
                metadata={'plan': plan, 'name': name, 'email': email},
                success_url=request.url_root + f'dashboard?access={{CHECKOUT_SESSION_ID}}&plan={plan}&name={name}&email={email}',
                cancel_url=request.url_root + 'uae',
            )
            return jsonify({"success": True, "checkout_url": session.url})
        else:
            # Stripe key not set yet — return empty so frontend falls back
            return jsonify({"success": False, "checkout_url": None})

    except Exception as e:
        logger.error(f"Stripe checkout error: {e}")
        return jsonify({"success": False, "checkout_url": None})

@app.route("/stripe-payment", methods=["POST"])
def stripe_payment():
    """Handle Stripe payment session creation"""
    try:
        data = request.get_json()
        plan = data.get('plan', 'starter')
        name = data.get('name', '')
        email = data.get('email', '')
        amount = data.get('amount', 3000)  # Amount in paise
        currency = data.get('currency', 'inr')
        
        if not plan or not name or not email:
            return jsonify({"error": "Missing required fields"}), 400
        
        # In production, use actual Stripe API
        # For now, simulate session creation
        session_id = f"cs_test_{datetime.now().strftime('%Y%m%d_%H%M%S')}_{hash(email + plan)}"
        
        # Create checkout session details
        session_data = {
            "session_id": session_id,
            "publishable_key": "pk_test_51234567890abcdef",  # Test key
            "checkout_url": f"https://checkout.stripe.com/pay/{session_id}",
            "amount": amount,
            "currency": currency,
            "plan": plan,
            "customer_email": email,
            "customer_name": name,
            "success_url": f"{request.url_root}payment-success?session_id={session_id}",
            "cancel_url": f"{request.url_root}payment-cancelled"
        }
        
        logger.info(f"Stripe session created: {session_id} for {email}")
        
        return jsonify({
            "success": True,
            "session_id": session_id,
            "checkout_url": session_data["checkout_url"],
            "publishable_key": session_data["publishable_key"]
        })
        
    except Exception as e:
        logger.error(f"Error creating Stripe session: {e}")
        return jsonify({"error": "Failed to create payment session"}), 500

@app.route("/stripe-webhook", methods=["POST"])
def stripe_webhook():
    """Handle Stripe webhook events"""
    try:
        # Verify webhook signature in production
        event = request.get_json()
        
        if event['type'] == 'checkout.session.completed':
            session = event['data']['object']
            email = session['customer_details']['email']
            plan = session['metadata'].get('plan', 'starter')
            name = email.split('@')[0] if email else 'Stripe Client'
            _, access_code = create_client(name, email, '', plan, days_valid=30)
            logger.info(f"Stripe payment completed: {session['id']} - Access: {access_code}")
            return jsonify({"status": "success", "access_code": access_code})
        
        return jsonify({"status": "received"})
        
    except Exception as e:
        logger.error(f"Stripe webhook error: {e}")
        return jsonify({"error": "Webhook processing failed"}), 500

@app.route("/payment-cancelled", methods=["GET", "POST"])
def payment_cancelled():
    """Handle cancelled payments"""
    return jsonify({
        "success": False,
        "message": "Payment was cancelled",
        "redirect_url": "/"
    })

@app.route("/paypal-ipn", methods=["POST"])
def paypal_ipn():
    """Handle PayPal Instant Payment Notification"""
    try:
        # In production, verify IPN data with PayPal
        logger.info("PayPal IPN received")
        return jsonify({"status": "verified"}), 200
    except Exception as e:
        logger.error(f"Error processing PayPal IPN: {e}")
        return jsonify({"status": "error"}), 500

@app.route("/chat", methods=["POST"])
def chat():
    data = request.get_json(silent=True)
    if not data:
        logger.warning("Received request with no JSON body")
        return jsonify({"error": "Request body must be JSON"}), 400

    user_message = data.get("message", "").strip()
    if not user_message:
        logger.warning("Received request with missing or empty 'message' field")
        return jsonify({"error": "Field 'message' is required and cannot be empty"}), 400

    access_code = data.get("access_code", "").strip()
    client = None
    if access_code:
        client = get_client_by_access_code(access_code)
        if not client:
            return jsonify({"error": "Invalid access code. Please check your embed code."}), 403

    language = data.get("language") or detect_language(user_message)
    timezone = data.get("timezone")
    history = data.get("history", [])
    region = (data.get("region") or "").strip().lower() or None
    if not region:
        cc = _detect_country_code()
        region_from_cc = {
            "IN": "india", "AE": "uae", "QA": "qatar", "KW": "kuwait",
            "BH": "bahrain", "OM": "oman", "BD": "bangladesh", "LK": "srilanka",
            "NP": "nepal", "CN": "china", "BT": "bhutan", "MM": "myanmar", "ID": "indonesia",
            "PK": "pakistan", "JO": "jordan", "IQ": "iraq", "LB": "lebanon",
            "PH": "philippines", "VN": "vietnam", "MY": "malaysia", "TH": "thailand",
            "NG": "nigeria", "KE": "kenya", "ZA": "southafrica", "EG": "egypt",
            "GH": "ghana", "UG": "uganda", "TZ": "tanzania", "RW": "rwanda",
            "MX": "mexico", "BR": "brazil", "CO": "colombia", "AR": "argentina",
            "PE": "peru", "CL": "chile",
        }
        if cc and cc in region_from_cc:
            region = region_from_cc[cc]
    context = ""
    if client:
        context = get_client_context(client['id'], user_message)
        increment_message_count(access_code)
        logger.info("Client chat: %s (id=%d, lang=%s)", access_code, client['id'], language)
    else:
        context = get_relevant_context(user_message)
    logger.info("Sending message to Claude (length=%d chars, lang=%s, tz=%s, history=%d, context=%d, region=%s)", len(user_message), language, timezone, len(history), len(context), region)

    try:
        reply = get_ai_reply(user_message, language, timezone, history, context, region)
    except anthropic.AuthenticationError as e:
        logger.error("Anthropic authentication failed — check ANTHROPIC_API_KEY: %s", e)
        return jsonify({"error": "API authentication failed. The server API key may be invalid or missing."}), 500
    except anthropic.RateLimitError as e:
        logger.error("Anthropic rate limit exceeded: %s", e)
        return jsonify({"error": "Rate limit exceeded. Please wait a moment and try again."}), 429
    except anthropic.APIStatusError as e:
        logger.error("Anthropic API returned status %s: %s", e.status_code, e.message)
        return jsonify({"error": f"Anthropic API error (status {e.status_code}): {e.message}"}), 502
    except anthropic.APIConnectionError as e:
        logger.error("Could not connect to Anthropic API: %s", e)
        return jsonify({"error": "Could not reach the Anthropic API. Check network connectivity."}), 502
    except Exception as e:
        logger.exception("Unexpected error while calling Anthropic API: %s", e)
        return jsonify({"error": f"Unexpected server error: {str(e)}"}), 500

    logger.info("Successfully received reply (length=%d chars)", len(reply))
    return jsonify({"reply": reply, "language": language})

@app.route("/upload", methods=["POST"])
def upload_document():
    if 'file' not in request.files:
        return jsonify({"error": "No file provided"}), 400
    file = request.files['file']
    if file.filename == '':
        return jsonify({"error": "No file selected"}), 400
    allowed_extensions = {'.pdf', '.docx', '.doc', '.txt', '.md', '.csv', '.json'}
    ext = os.path.splitext(file.filename.lower())[1]
    if ext not in allowed_extensions:
        return jsonify({"error": f"Unsupported file type: {ext}. Allowed: {', '.join(allowed_extensions)}"}), 400

    try:
        text = extract_text_from_file(file, file.filename)
    except Exception as e:
        logger.error("Extraction failed: %s", e)
        return jsonify({"error": "Failed to extract text from file"}), 500

    if not text.strip():
        return jsonify({"error": "No text could be extracted from the file"}), 400

    chunk_count = add_document_to_knowledge_base(file.filename, text)
    preview = text[:300].replace('\n', ' ')
    return jsonify({
        "success": True,
        "filename": file.filename,
        "chunks": chunk_count,
        "preview": preview
    })

@app.route("/speak", methods=["POST"])
def speak():
    data = request.get_json(silent=True)
    text = data.get("text", "") if data else ""
    language = data.get("language", "en") if data else "en"
    if not text or not elevenlabs_key:
        return jsonify({"error": "Missing text or API key"}), 400

    voice_id = ELEVENLABS_VOICES.get(language, ELEVENLABS_VOICES["en"])
    url = f"https://api.elevenlabs.io/v1/text-to-speech/{voice_id}"
    headers = {
        "Accept": "audio/mpeg",
        "Content-Type": "application/json",
        "xi-api-key": elevenlabs_key
    }
    payload = {
        "text": text[:4000],
        "model_id": "eleven_multilingual_v2",
        "voice_settings": {
            "stability": 0.5,
            "similarity_boost": 0.5
        }
    }

    try:
        resp = requests.post(url, json=payload, headers=headers, timeout=60)
        if resp.status_code == 401 or resp.status_code == 429:
            logger.warning("ElevenLabs quota exceeded or unauthorized - client will use browser TTS")
            return jsonify({"error": "quota_exceeded"}), 503
        if resp.status_code != 200:
            logger.error("ElevenLabs error: status=%s, body=%s", resp.status_code, resp.text[:500])
            return jsonify({"error": "ElevenLabs error"}), 503
        return Response(resp.content, mimetype="audio/mpeg")
    except Exception as e:
        logger.error("ElevenLabs error: %s", e)
        return jsonify({"error": "Failed to generate audio"}), 503

EDIT_UPLOAD_DIR = os.environ.get("EDIT_UPLOAD_DIR", "/tmp/botifyai_edits")
EDIT_OUTPUT_DIR = os.environ.get("EDIT_OUTPUT_DIR", "/tmp/botifyai_edits/out")
EDIT_SESSION_DIR = os.environ.get("EDIT_SESSION_DIR", "/tmp/botifyai_edits/sessions")
os.makedirs(EDIT_UPLOAD_DIR, exist_ok=True)
os.makedirs(EDIT_OUTPUT_DIR, exist_ok=True)
os.makedirs(EDIT_SESSION_DIR, exist_ok=True)

EDIT_SESSION_TTL_SECONDS = 30 * 60  # 30 minutes


def _session_dir(session_id: str) -> str:
    """Safe path for a session's working files. Rejects traversal."""
    if not re.fullmatch(r"[A-Za-z0-9_-]{6,64}", session_id or ""):
        raise ValueError("Invalid session_id")
    return os.path.join(EDIT_SESSION_DIR, session_id)


def _session_current_file(session_id: str) -> str | None:
    """Return the path to the session's most recent working file, if any and
    not expired."""
    try:
        sdir = _session_dir(session_id)
    except ValueError:
        return None
    if not os.path.isdir(sdir):
        return None
    candidates = [
        os.path.join(sdir, name)
        for name in os.listdir(sdir)
        if name.startswith("current.")
    ]
    if not candidates:
        return None
    path = candidates[0]
    age = time.time() - os.path.getmtime(path)
    if age > EDIT_SESSION_TTL_SECONDS:
        # Expired — clean up
        try:
            shutil.rmtree(sdir, ignore_errors=True)
        except OSError:
            pass
        return None
    return path


def _session_set_current(session_id: str, source_path: str) -> str:
    """Copy source into the session's current-file slot; return new path."""
    sdir = _session_dir(session_id)
    os.makedirs(sdir, exist_ok=True)
    # Clear any prior current.*
    for name in os.listdir(sdir):
        if name.startswith("current."):
            try:
                os.remove(os.path.join(sdir, name))
            except OSError:
                pass
    ext = os.path.splitext(source_path)[1] or ".bin"
    dst = os.path.join(sdir, f"current{ext}")
    shutil.copy(source_path, dst)
    return dst


@app.route("/edit-media/ops-schema", methods=["GET"])
def edit_media_schema():
    """Clients (LumaEdit app) fetch this to know what ops are supported."""
    return Response(SUPPORTED_OPS_HINT, mimetype="application/json")


@app.route("/edit-media", methods=["POST"])
def edit_media_endpoint():
    """Multipart upload for chat-style iterative editing.

    Form fields:
      file:         (required on first turn) the source photo/video
      session_id:   (optional) client-generated stable id; if the server has
                    a working file for this id, the ops apply to it instead
                    of the uploaded file. On subsequent turns, sending only
                    session_id + instruction is enough.
      instruction:  natural-language description; Claude parses to ops
      ops:          JSON op list (bypasses the parser; for manual UI)

    Response:
      { success, session_id, kind, ops_applied, summary,
        download_url, filename }
    """
    session_id = (request.form.get("session_id") or "").strip()
    if session_id:
        try:
            _session_dir(session_id)  # validates format
        except ValueError as e:
            return jsonify({"error": str(e)}), 400

    # Decide input path: session's current file (subsequent turn) OR upload.
    input_path: str | None = None
    uploaded_temp: str | None = None
    session_input = _session_current_file(session_id) if session_id else None

    if session_input:
        input_path = session_input
    else:
        if "file" not in request.files or not request.files["file"].filename:
            return jsonify({
                "error": "First turn requires a file upload (or a valid session_id "
                         "with a previously uploaded file that hasn't expired)."
            }), 400
        file = request.files["file"]
        try:
            kind_from_path(file.filename)
        except MediaEditorError as e:
            return jsonify({"error": str(e)}), 400
        safe_name = re.sub(r"[^A-Za-z0-9._-]", "_", file.filename)
        uploaded_temp = os.path.join(EDIT_UPLOAD_DIR, f"{uuid.uuid4().hex}_{safe_name}")
        file.save(uploaded_temp)
        input_path = uploaded_temp

    kind = kind_from_path(input_path)

    # Parse instruction or accept explicit ops.
    ops: list = []
    summary = ""
    ops_json = request.form.get("ops")
    instruction = (request.form.get("instruction") or "").strip()

    def _cleanup_upload() -> None:
        if uploaded_temp:
            try:
                os.remove(uploaded_temp)
            except OSError:
                pass

    if ops_json:
        try:
            ops = json.loads(ops_json)
            if not isinstance(ops, list):
                raise ValueError("ops must be a JSON list")
        except Exception as e:
            _cleanup_upload()
            return jsonify({"error": f"Invalid ops JSON: {e}"}), 400
        summary = f"Applied {len(ops)} operation(s)"
    elif instruction:
        parsed = parse_edit_instruction(get_anthropic_client(), kind, instruction)
        ops = parsed["ops"]
        summary = parsed["summary"]
        if not ops:
            _cleanup_upload()
            return jsonify({
                "error": summary or "No editable operations found in instruction",
                "instruction": instruction,
            }), 400
    else:
        _cleanup_upload()
        return jsonify({"error": "Provide either 'instruction' or 'ops' form field"}), 400

    try:
        output_path, out_kind = edit_media(input_path, ops, EDIT_OUTPUT_DIR)
    except MediaEditorError as e:
        _cleanup_upload()
        return jsonify({"error": str(e), "ops": ops}), 400
    except Exception as e:
        _cleanup_upload()
        logger.exception("edit_media crashed: %s", e)
        return jsonify({"error": f"Internal edit error: {e}"}), 500

    # Persist output as the session's new current file (create session if none).
    if not session_id:
        session_id = uuid.uuid4().hex
    try:
        _session_set_current(session_id, output_path)
    except Exception as e:
        logger.warning("Could not persist session state: %s", e)

    _cleanup_upload()

    out_name = os.path.basename(output_path)
    return jsonify({
        "success": True,
        "session_id": session_id,
        "kind": out_kind,
        "input_kind": kind,
        "ops_applied": ops,
        "summary": summary,
        "download_url": f"/edit-media/download/{out_name}",
        "filename": out_name,
    })


@app.route("/edit-media/session/<session_id>", methods=["DELETE"])
def edit_media_reset_session(session_id: str):
    """Clear a session's working state (client's 'Reset' button)."""
    try:
        sdir = _session_dir(session_id)
    except ValueError as e:
        return jsonify({"error": str(e)}), 400
    if os.path.isdir(sdir):
        shutil.rmtree(sdir, ignore_errors=True)
    return jsonify({"success": True, "session_id": session_id})


@app.route("/transcribe-media", methods=["POST"])
def transcribe_media():
    """Extract audio from an uploaded video/audio file, run Whisper on it via
    Replicate, and return the resulting .srt as a downloadable file. Cost:
    ~$0.006 per minute of audio.
    """
    if "file" not in request.files or not request.files["file"].filename:
        return jsonify({"error": "No file provided (field name: 'file')"}), 400
    file = request.files["file"]
    try:
        kind_from_path(file.filename)
    except MediaEditorError as e:
        return jsonify({"error": str(e)}), 400
    safe_name = re.sub(r"[^A-Za-z0-9._-]", "_", file.filename)
    src_path = os.path.join(EDIT_UPLOAD_DIR, f"{uuid.uuid4().hex}_{safe_name}")
    file.save(src_path)

    # Extract mono 16 kHz audio (Whisper's preferred input) to keep upload small.
    audio_path = os.path.join(EDIT_UPLOAD_DIR, f"{uuid.uuid4().hex}.wav")
    try:
        proc = subprocess.run(
            [
                "ffmpeg", "-y", "-i", src_path,
                "-vn", "-ac", "1", "-ar", "16000", "-c:a", "pcm_s16le",
                audio_path,
            ],
            capture_output=True,
            text=True,
            timeout=120,
        )
        if proc.returncode != 0:
            return jsonify({"error": f"Audio extract failed: {proc.stderr[-400:]}"}), 500

        try:
            srt_path, words_path = ai_transcribe_words(audio_path, EDIT_OUTPUT_DIR)
        except ReplicateUnavailable as e:
            return jsonify({"error": str(e)}), 503
        except Exception as e:
            logger.exception("whisperx failed")
            return jsonify({"error": f"Transcription failed: {e}"}), 500
        return jsonify({
            "success": True,
            "download_url": f"/edit-media/download/{os.path.basename(srt_path)}",
            "filename": os.path.basename(srt_path),
            # Sidecar with per-word timings from WhisperX alignment. The
            # editor uses this to build karaoke ASS where each word lights
            # up on its true onset. Optional — legacy clients ignoring
            # this field still get a working SRT.
            "words_url": f"/edit-media/download/{os.path.basename(words_path)}",
            "words_filename": os.path.basename(words_path),
        })
    finally:
        for p in (src_path, audio_path):
            try: os.remove(p)
            except OSError: pass


@app.route("/detect-beats", methods=["POST"])
def detect_beats_endpoint():
    """Onset / beat detection on an uploaded audio (or video) file. Uses a
    pure-numpy RMS-envelope + rising-edge threshold, so no extra
    dependencies beyond what's already in requirements.txt.

    Not as accurate as librosa's beat_track for musical BPM tracking,
    but plenty good for editor-side "split on rhythm hits" — which is
    what CapCut's Beat Sync ships in practice.

    Form fields:
      file:      required — the music (mp3/wav/m4a) OR any video file;
                  audio is auto-extracted via ffmpeg.
      min_gap:   optional — minimum seconds between adjacent beats
                  (default 0.20 ≈ 300 BPM ceiling).
      strength:  optional — 0..1, higher = fewer beats. Default 0.5.

    Response: { success, beats: [seconds…], count }
    """
    if "file" not in request.files or not request.files["file"].filename:
        return jsonify({"error": "No file provided (field name: 'file')"}), 400
    file = request.files["file"]
    min_gap = float(request.form.get("min_gap") or 0.20)
    strength = float(request.form.get("strength") or 0.5)

    import numpy as np
    safe_name = re.sub(r"[^A-Za-z0-9._-]", "_", file.filename)
    src_path = os.path.join(EDIT_UPLOAD_DIR, f"{uuid.uuid4().hex}_{safe_name}")
    file.save(src_path)
    try:
        sr = 22050
        # Decode audio to raw 16-bit mono @ sr Hz via ffmpeg pipe.
        proc = subprocess.run(
            ["ffmpeg", "-y", "-i", src_path, "-vn", "-ac", "1", "-ar", str(sr),
             "-f", "s16le", "-loglevel", "error", "pipe:1"],
            capture_output=True, timeout=120,
        )
        if proc.returncode != 0:
            return jsonify({"error": f"Audio decode failed: {proc.stderr[:400]}"}), 500
        pcm = np.frombuffer(proc.stdout, dtype=np.int16).astype(np.float32) / 32768.0
        if pcm.size < sr:
            return jsonify({"error": "Audio too short for beat detection"}), 400

        # RMS envelope — hop=512 samples ≈ 23 ms per frame.
        hop = 512
        n = (pcm.size - hop) // hop
        env = np.empty(n, dtype=np.float32)
        for i in range(n):
            block = pcm[i * hop:i * hop + hop]
            env[i] = float(np.sqrt(np.mean(block * block)))
        # Rising-edge onset function: positive differences of the envelope.
        diff = np.diff(env)
        diff = np.maximum(diff, 0)
        # Adaptive threshold — median + k * std. strength maps to k in [1..3].
        k = 1.0 + 2.0 * max(0.0, min(1.0, strength))
        thr = float(np.median(diff)) + k * float(np.std(diff))
        peaks = np.where(diff > thr)[0]

        # De-duplicate: enforce min_gap seconds between accepted beats.
        beats = []
        min_gap_frames = max(1, int(min_gap * sr / hop))
        last = -min_gap_frames * 2
        for p in peaks:
            if p - last >= min_gap_frames:
                beats.append(round(float(p * hop / sr), 3))
                last = p
        return jsonify({"success": True, "beats": beats, "count": len(beats)})
    finally:
        try:
            os.remove(src_path)
        except OSError:
            pass


@app.route("/generate-voiceover", methods=["POST"])
def generate_voiceover_endpoint():
    """Generate an AI voice-over audio track for the media editor via
    ElevenLabs. Takes a script + optional language; returns a mp3
    file URL just like /transcribe-media so the editor can attach it
    as a voiceOver MusicTrack.

    Form fields:
      text:     (required) the script to speak (<= 3000 chars trimmed)
      language: (optional) full language name — defaults to English
      stability: (optional) 0..1, ElevenLabs stability. Default 0.5
      similarity: (optional) 0..1, similarity_boost. Default 0.75

    Response: { success, download_url, filename }
    Errors:
      400 — missing text
      503 — ElevenLabs key missing / quota exceeded / upstream error
    """
    text = (request.form.get("text") or "").strip()
    if not text:
        return jsonify({"error": "Provide 'text' — the script to speak"}), 400
    if not elevenlabs_key:
        return jsonify({"error": "ELEVENLABS_API_KEY not set on server"}), 503

    language = (request.form.get("language") or "English").strip()
    stability = float(request.form.get("stability") or 0.5)
    similarity = float(request.form.get("similarity") or 0.75)

    # Map full language name → voice id via the existing voice table;
    # unknown languages fall through to the default English voice.
    lang_key = "en"
    for k, name in LANGUAGE_NAMES.items():
        if name.lower() == language.lower():
            lang_key = k
            break
    voice_id = ELEVENLABS_VOICES.get(lang_key, ELEVENLABS_VOICES["en"])

    url = f"https://api.elevenlabs.io/v1/text-to-speech/{voice_id}"
    headers = {
        "Accept": "audio/mpeg",
        "Content-Type": "application/json",
        "xi-api-key": elevenlabs_key,
    }
    payload = {
        "text": text[:3000],
        "model_id": "eleven_multilingual_v2",
        "voice_settings": {
            "stability": max(0.0, min(1.0, stability)),
            "similarity_boost": max(0.0, min(1.0, similarity)),
        },
    }
    try:
        resp = requests.post(url, json=payload, headers=headers, timeout=120)
    except Exception as e:
        logger.exception("ElevenLabs request failed")
        return jsonify({"error": f"ElevenLabs upstream error: {e}"}), 503
    if resp.status_code in (401, 429):
        return jsonify({"error": "ElevenLabs quota exceeded or key rejected"}), 503
    if resp.status_code != 200:
        return jsonify({"error": f"ElevenLabs returned {resp.status_code}"}), 503

    out_path = os.path.join(EDIT_OUTPUT_DIR, f"voiceover_{uuid.uuid4().hex}.mp3")
    with open(out_path, "wb") as f:
        f.write(resp.content)
    return jsonify({
        "success": True,
        "download_url": f"/edit-media/download/{os.path.basename(out_path)}",
        "filename": os.path.basename(out_path),
    })


@app.route("/generate-caption", methods=["POST"])
def generate_caption_endpoint():
    """Generate a punchy social-media caption + hashtags for a short-form
    video via Claude Haiku. One-shot text generation — no source file.

    Form fields:
      topic:     (required) short description of the video's subject
      platform:  (optional) 'instagram' | 'tiktok' | 'youtube_shorts' |
                  'twitter' | 'linkedin'. Default 'instagram'.
      language:  (optional) full language name for the caption
                  (e.g. 'English', 'Hindi', 'Spanish'). Default English.
      tone:      (optional) 'casual' | 'professional' | 'funny' |
                  'inspirational'. Default 'casual'.

    Response: { success, caption, hashtags: [str] }
    """
    topic = (request.form.get("topic") or "").strip()
    if not topic:
        return jsonify({"error": "Provide 'topic' — short description of the video"}), 400
    platform = (request.form.get("platform") or "instagram").strip().lower()
    language = (request.form.get("language") or "English").strip()
    tone = (request.form.get("tone") or "casual").strip().lower()

    client = get_anthropic_client()
    if not client:
        return jsonify({"error": "AI unavailable"}), 503

    # Platform-specific caption length + hashtag count guidance.
    guidance = {
        "instagram": ("1-2 short sentences with 1-2 emojis", 8, 12),
        "tiktok": ("one punchy line, ≤80 chars, hooky, 1 emoji ok", 4, 6),
        "youtube_shorts": ("one hook line + one CTA", 5, 8),
        "twitter": ("≤240 chars, minimal hashtags", 2, 4),
        "linkedin": ("2-3 sentences, professional, minimal emoji", 3, 5),
    }.get(platform, ("1-2 short sentences with 1-2 emojis", 8, 12))
    fmt, hmin, hmax = guidance

    system = (
        f"You write social-media captions in the user's chosen language "
        f"({language}). Return ONLY valid JSON, no prose, no code fences:\n"
        f'{{"caption": "<caption text>", "hashtags": ["#tag1", "#tag2", ...]}}\n\n'
        f"Constraints:\n"
        f"- Platform: {platform}. Caption style: {fmt}.\n"
        f"- Tone: {tone}.\n"
        f"- Include {hmin}-{hmax} relevant hashtags (English hashtags stay in "
        f"English even when the caption is in another language — that's the "
        f"platform norm).\n"
        f"- No repeating the user's topic verbatim; add hook / emotion."
    )
    try:
        resp = client.messages.create(
            model="claude-haiku-4-5-20251001",
            max_tokens=500,
            system=system,
            messages=[{"role": "user", "content": f"Topic: {topic}"}],
        )
        text = "".join(
            b.text for b in resp.content if getattr(b, "type", "") == "text"
        ).strip()
        text = re.sub(r"^```(?:json)?\s*", "", text)
        text = re.sub(r"\s*```$", "", text)
        parsed = json.loads(text)
    except Exception as e:
        logger.exception("caption generation failed")
        return jsonify({"error": f"Generation failed: {e}"}), 500

    return jsonify({
        "success": True,
        "platform": platform,
        "language": language,
        "caption": (parsed.get("caption") or "").strip(),
        "hashtags": [str(h).strip() for h in (parsed.get("hashtags") or [])],
    })


@app.route("/translate-captions", methods=["POST"])
def translate_captions_endpoint():
    """Translate an SRT caption file to a target language via Claude Haiku.
    Preserves timestamps exactly — only the text lines are translated.

    Form fields:
      file:            (required) the source .srt file
      target_language: (required) full language name, e.g. "Spanish",
                        "Hindi", "Arabic". Free-form; Claude handles it.
    """
    if "file" not in request.files or not request.files["file"].filename:
        return jsonify({"error": "No file provided (field name: 'file')"}), 400
    target = (request.form.get("target_language") or "").strip()
    if not target:
        return jsonify({"error": "Provide 'target_language' (e.g. Spanish, Hindi)"}), 400

    src_file = request.files["file"]
    srt_text = src_file.read().decode("utf-8", errors="replace")
    if not srt_text.strip():
        return jsonify({"error": "SRT file is empty"}), 400

    client = get_anthropic_client()
    if not client:
        return jsonify({"error": "Translator unavailable"}), 503

    system = (
        "You translate SRT subtitle files. Preserve EVERY timestamp line and "
        "the numeric index lines EXACTLY. Only translate the spoken-text lines "
        "into the target language. Keep line breaks. Output ONLY the translated "
        "SRT, no prose, no code fences."
    )
    try:
        resp = client.messages.create(
            model="claude-haiku-4-5-20251001",
            max_tokens=6000,
            system=system,
            messages=[{
                "role": "user",
                "content": f"Target language: {target}\n\n{srt_text}",
            }],
        )
        translated = "".join(
            b.text for b in resp.content if getattr(b, "type", "") == "text"
        ).strip()
    except Exception as e:
        logger.exception("caption translation failed")
        return jsonify({"error": f"Translation failed: {e}"}), 500

    os.makedirs(EDIT_OUTPUT_DIR, exist_ok=True)
    out_name = f"captions_{uuid.uuid4().hex}_{target.replace(' ', '_')}.srt"
    out_path = os.path.join(EDIT_OUTPUT_DIR, out_name)
    with open(out_path, "w", encoding="utf-8") as f:
        f.write(translated)

    return jsonify({
        "success": True,
        "download_url": f"/edit-media/download/{out_name}",
        "filename": out_name,
        "target_language": target,
    })


@app.route("/generate-image", methods=["POST"])
def generate_image_endpoint():
    """Generate an image from a text prompt via Flux Schnell on Replicate
    (~$0.003 per image — cheap). One-shot; not tied to a session.

    Form fields:
      prompt:       (required) text description
      aspect_ratio: (optional) '1:1' | '9:16' | '16:9' | '4:5'. Default '9:16'.
    """
    from replicate_ops import ai_text_to_image
    prompt = (request.form.get("prompt") or "").strip()
    if not prompt:
        return jsonify({"error": "Provide a 'prompt'"}), 400
    aspect = (request.form.get("aspect_ratio") or "9:16").strip()
    try:
        out_path = ai_text_to_image(prompt, {"aspect_ratio": aspect}, EDIT_OUTPUT_DIR)
    except ReplicateUnavailable as e:
        return jsonify({"error": str(e)}), 503
    except Exception as e:
        logger.exception("text-to-image failed")
        return jsonify({"error": f"Generation failed: {e}"}), 500
    return jsonify({
        "success": True,
        "kind": "image",
        "download_url": f"/edit-media/download/{os.path.basename(out_path)}",
        "filename": os.path.basename(out_path),
    })


@app.route("/generate-video", methods=["POST"])
def generate_video_endpoint():
    """Generate a ~6 second video from a text prompt via MiniMax Hailuo on
    Replicate. Cost: ~$0.30 per call. Not tied to a session — one-shot.

    Form fields:
      prompt: str (required)
    """
    from replicate_ops import ai_text_to_video
    prompt = (request.form.get("prompt") or "").strip()
    if not prompt:
        return jsonify({"error": "Provide a 'prompt'"}), 400
    try:
        out_path = ai_text_to_video(prompt, {}, EDIT_OUTPUT_DIR)
    except ReplicateUnavailable as e:
        return jsonify({"error": str(e)}), 503
    except Exception as e:
        logger.exception("text-to-video failed")
        return jsonify({"error": f"Generation failed: {e}"}), 500
    return jsonify({
        "success": True,
        "kind": "video",
        "download_url": f"/edit-media/download/{os.path.basename(out_path)}",
        "filename": os.path.basename(out_path),
    })


@app.route("/generate-music", methods=["POST"])
def generate_music_endpoint():
    """Generate an original music clip from a text prompt via Meta
    MusicGen on Replicate. Copyright-safe (AI-composed). Cost ~$0.02
    per call for 15s. Not tied to a session — one-shot.

    Form fields:
      prompt: str (required, e.g. "upbeat bollywood beat with tabla")
      duration: int (optional, 4-30 seconds, default 15)
    """
    from replicate_ops import ai_generate_music
    prompt = (request.form.get("prompt") or "").strip()
    if not prompt:
        return jsonify({"error": "Provide a 'prompt'"}), 400
    try:
        duration = int(request.form.get("duration", "15"))
    except ValueError:
        duration = 15
    try:
        out_path = ai_generate_music(
            prompt, {"duration": duration}, EDIT_OUTPUT_DIR)
    except ReplicateUnavailable as e:
        return jsonify({"error": str(e)}), 503
    except Exception as e:
        logger.exception("music generation failed")
        return jsonify({"error": f"Generation failed: {e}"}), 500
    return jsonify({
        "success": True,
        "kind": "music",
        "download_url": f"/edit-media/download/{os.path.basename(out_path)}",
        "filename": os.path.basename(out_path),
    })


@app.route("/edit-media/download/<path:name>", methods=["GET"])
def edit_media_download(name: str):
    # basic path traversal guard
    if "/" in name or ".." in name:
        return jsonify({"error": "Invalid filename"}), 400
    return send_from_directory(EDIT_OUTPUT_DIR, name, as_attachment=True)


@app.route("/whatsapp", methods=["POST"])
def whatsapp():
    sender = request.form.get("From", "")
    message = request.form.get("Body", "").strip()

    if not message or is_group_message(sender):
        return str(MessagingResponse())

    language = detect_language(message)
    reply = get_ai_reply(message, language)

    if len(reply) > 1500:
        reply = reply[:1497] + "..."

    send_whatsapp_reply(sender, reply)
    return str(MessagingResponse())

if __name__ == "__main__":
    app.run(host="0.0.0.0", port=port, debug=True)

